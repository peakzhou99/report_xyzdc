import json
import time
from analyze.private_apply import PrivateApply
from data.doc_data.split_doc import SplitDoc
from data.db_data.yjt_db_data import YjtDBDataHandler
from utils.table_util import create_table, create_table_by_oxml
from utils.doc_util import *
from utils.file_util import *
from utils.string_util import *
from utils.log_utils import get_logger
from utils.get_pf_api import CreditApprovalFetcher
from docx import Document
from structured.gen_dynamic_sections import ComprehensiveReportGenerator
logger = get_logger()


class UrbanReport():
    def __init__(self, request_data, company_name, root_data_dir):
        # 路径信息
        self.root_data_dir = root_data_dir
        self.request_data = request_data
        self.external_data_dir = os.path.join(root_data_dir, f"{company_name}/external_data")
        self.company_name = company_name
        self.prospectus_doc_path = None
        self.start_time = None

        # 解析的数据信息
        self.table_datas = {}
        self.table_oxmls = {}
        self.article_datas = {}
        self.company_info = {}
        self.fuzhai_total = None
        self.element = "【债券要素提取错误】"
        # 风险点总结
        self.risk = []
        self.risk_evaluate = []

        # 实例初始化
        self.private_apply_instance = PrivateApply()

        # 查找获取路径中的 募集说明书、调查报告
        for root, dirs, files in os.walk(self.root_data_dir):
            for file in files:
                if "募集说明书" in file and file.endswith(".docx"):
                    self.prospectus_doc_path = os.path.join(root, file)

        if not self.prospectus_doc_path:
            raise Exception("文件不符合要求。")

        self.fetcher = CreditApprovalFetcher()

        self.comprehensive_generator = ComprehensiveReportGenerator(company_name)

    def prepare_data(self):

        # 募集说明书
        prospectusDoc = SplitDoc(self.prospectus_doc_path, self.company_name)
        # 解析文本信息
        prospectusDoc.extract_paragraphs()
        self.company_info = prospectusDoc.company_info

        # 解析表格信息
        prospectusDoc.extract_tables()
        self.fuzhai_total = prospectusDoc.fuzhai_total

        # 预警通数据
        yjt_handler = YjtDBDataHandler(self.company_name, self.root_data_dir)
        yjt_table_datas = yjt_handler.get_yjt_data()

        # 文本数据融合（调查文本与募集文本融合）
        self.article_datas = prospectusDoc.article_datas
        # 发行平台级别
        self.article_datas["发行平台级别"] = ""
        for yjt_article_name, yjt_article_data in yjt_handler.article_datas.items():
            if self.article_datas.get(yjt_article_name):
                self.article_datas[yjt_article_name].append(yjt_article_data)
            else:
                self.article_datas[yjt_article_name] = yjt_article_data

        # 表格数据融合（外部报表与募集报表融合）,应收，其他应收，资产负债，现金流
        _table_datas = prospectusDoc.table_datas
        for yjt_data in yjt_table_datas:
            if "发债平台情况" == yjt_data['name'] and len(yjt_data['data']) >= 2:
                platform_level = f"发行人为{yjt_data['data'][1][1]}{yjt_data['data'][1][2]}平台"
                self.article_datas["发行平台级别"] = platform_level
            if _table_datas.__contains__(yjt_data['name']):
                _table_datas[yjt_data['name']].append(yjt_data)
            else:
                _table_datas[yjt_data['name']] = [yjt_data]
        # 筛选最新表格数据
        for tab_title, tab_data_list in _table_datas.items():
            last_tab_data = {}
            last_date = '197001'
            for tab_data in tab_data_list:
                if tab_data['date'] is None and len(tab_data_list) == 1:
                    last_date = ''
                    last_tab_data = tab_data
                elif tab_data['date'] is not None and tab_data['date'] > last_date:
                    last_date = tab_data['date']
                    last_tab_data = tab_data
            self.table_datas[tab_title] = last_tab_data
            # 从有息负债表中提取非标融资
            self._extract_non_standard_financing()

        # logger.debug("表格数据：",self.table_datas)
        # 筛选最新oxml数据
        for oxml_title, oxml_data_list in prospectusDoc.table_oxmls.items():
            last_oxml_data = {}
            last_date = '197001'
            for oxml_data in oxml_data_list:
                if oxml_data['date'] is None and len(oxml_data_list) == 1:
                    last_date = ''
                    last_oxml_data = oxml_data
                elif oxml_data['date'] is not None and oxml_data['date'] >= last_date:
                    last_date = oxml_data['date']
                    last_oxml_data = oxml_data
            self.table_oxmls[oxml_title] = last_oxml_data

        # 提取并计算风险信息
        if self.table_datas.get("资产负债表"):
            assets_total = self.risk_for_assets_liabilities(self.table_datas["资产负债表"])
        if self.table_datas.get("营业收入"):
            self.risk_for_operating_revenue(self.table_datas["营业收入"])
        if self.table_datas.get("有息负债"):
            self.risk_for_interest_bearing_debt(self.table_datas["有息负债"])
        if self.table_datas.get("受限资产"):
            self.risk_for_restricted_assets(self.table_datas["受限资产"], assets_total)

    def _extract_non_standard_financing(self):
        """从有息负债表中提取非标融资数据"""
        logger.debug("提取非标融资表格")

        # 获取有息负债-募集表
        interest_bearing_table = self.table_datas.get("有息负债-募集")
        if not interest_bearing_table:
            logger.debug("未找到有息负债表，无法提取非标融资")
            return

        table_data = interest_bearing_table.get("data", [])
        if not table_data:
            logger.debug("有息负债表数据为空")
            return

        # 将表格数据转换为字符串格式供LLM处理
        table_content = "\n".join(["\t".join([str(cell) for cell in row]) for row in table_data])

        extract_messages = [{
            "role": "user",
            "content": f"""/no_think
            你是财务表格分析专家，从提供的有息负债表格中提取仅属于"非标融资"子类别的行。非标融资指信托、融资租赁、私募债等非标准化融资项目。保留完整的表头行和所有相关数据行（包括空行），但在遇到"其他融资"、"合计"、"小计"或"地方专项债券转贷"等汇总或非具体非标融资行时停止提取。确保输出为JSON格式，表头作为第一行。
        
            【表格内容】
            {table_content}
        
            【输出格式】
            {{"sub_table": [["表头1", "表头2"], ["行1值1", "行1值2"], ...]}}
        
            如果没有相关内容，输出 {{"sub_table": []}}
            """
        }]

        try:
            result = self.private_apply_instance.generate(extract_messages)
            extracted = json.loads(result)
            sub_table = extracted.get("sub_table", [])

            if sub_table and len(sub_table) > 1:  # 确保有表头和数据
                # 工程化处理：如果提取出"其他融资"，从这一行及后面的数据直接去掉
                header = sub_table[0]
                data_rows = sub_table[1:]
                filtered_data_rows = []

                for row in data_rows:
                    # 检查是否包含停止关键词
                    if any(keyword in str(cell) for cell in row for keyword in
                           ["其他融资", "合计", "小计", "地方专项债券转贷"]):
                        break
                    filtered_data_rows.append(row)

                if filtered_data_rows:  # 确保过滤后仍有数据
                    sub_table = [header] + filtered_data_rows

                    # 创建非标融资表格数据
                    non_standard_table = {
                        "name": "非标融资",
                        "date": interest_bearing_table.get("date"),
                        "unit": interest_bearing_table.get("unit"),
                        "preceding_text": f"发行人{interest_bearing_table.get('date', '')}非标融资情况",
                        "data": sub_table
                    }

                    # 添加到表格数据集合中
                    self.table_datas["非标融资"] = non_standard_table
                    logger.debug(f"成功提取非标融资数据，共{len(filtered_data_rows)}行")
                else:
                    logger.debug("过滤后无有效非标融资数据")
            else:
                logger.debug("LLM未提取到非标融资数据")

        except Exception as e:
            logger.exception(f"提取非标融资失败: {e}")


    def risk_for_restricted_assets(self, table, assets_total):
        llm_prompt = [{
            "role": "user",
            "content": f"""/no_think
            你是金融文档分析专家，请从“受限资产表格”中提取关键信息，并以JSON格式输出：
            1. 提取受限资产总金额，如果没有则不提取；
            2. 根据表格标题和表头最新的时间序列，提取报告截至日期，格式如'X年X月末'或'X年末'，如果没有则不提取。
            3. 结果以JSON格式输出；
            【受限资产表格】
            {table['data']}
            【表格标题】
            {table.get('preceding_text', "")}
            【结果输出】
            {{"受限资产总金额":0.0,"报告截至日期":""}}
            """
        }]
        _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
        if "亿" not in table['unit']:
            _result_json = self.unit_conversion(_result_json, table['unit'])
        restricted_assets = float(_result_json["受限资产总金额"]) if _result_json["受限资产总金额"] else _result_json[
            "受限资产总金额"]
        restricted_assets_date = _result_json["报告截至日期"] or "报告期末"
        if restricted_assets and assets_total:
            proportion = restricted_assets / assets_total
            if proportion > 0.2:
                self.risk_evaluate.append(f"受限资产占总收入比例超{round(proportion * 10 + 1) * 10}%")
                self.risk.append(f"受限资产占比高：截至{restricted_assets_date}，发行人总资产{round(assets_total, 2)}亿元，其中受限资产{round(restricted_assets, 2)}亿元，受限资产占总收入比例超{round(proportion * 100, 2)}%。")

    def risk_for_interest_bearing_debt(self, table):
        try:
            table_data = table.get("data")
            if table_data[0][0] == "项目" and "合计" in table_data[-1][0]:
                change_flag = False
                res_table = []
                for index, row in enumerate(table_data):
                    if row[0] == "项目":
                        res_table.append(row)
                    elif row[0] == "银行贷款":
                        change_flag = True
                        break
                    else:
                        break
                if change_flag:
                    res_table.append(table_data[-1])
                    table["data"] = res_table
        except Exception as e:
            logger.exception(f"youxi table handle error:{e}")
        llm_prompt = [{
            "role": "user",
            "content": f"""/no_think
            你是金融文档分析专家，请从“有息负债表格”中提取关键信息，并以JSON格式输出：
            1. 提取短期债务（一年期以内）金额，直接取表格最后的合计数据，如果没有则不提取。
            2. 解析表格中是否有多个时间序列，如果有，就对各时间序列的数据做出划分。
            3. 提取有息负债总额，如果有多个值，必须取最新的一期。
            4. 根据表格标题或表格表头，提取统计时间，如果有多个时间序列，必须取最新的一期，格式如：'xxxx年xx月末'、'xxxx年末'，如果没有则不提取。
            5. 结果以JSON格式输出。
            【有息负债表格】
            {table['data']}
            【表格标题】
            {table.get('preceding_text', "")}
            【结果输出】
            {{"短期债务":0.0,"有息负债总额":0.0,"统计时间":""}}
            """
        }]
        _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
        if "亿" not in table['unit']:
            _result_json = self.unit_conversion(_result_json, table['unit'])
        short_term_debt = float(_result_json["短期债务"]) if _result_json["短期债务"] else _result_json["短期债务"]
        interest_bearing_debt = float(_result_json["有息负债总额"]) if _result_json["有息负债总额"] else _result_json[
            "有息负债总额"]
        debt_date = _result_json["统计时间"] or "报告期末"
        if short_term_debt and interest_bearing_debt:
            proportion = short_term_debt / interest_bearing_debt
            if proportion > 0.3:
                self.risk_evaluate.append(f"短期债务占有息负债比例超{round(proportion * 10 + 1) * 10}%")
                self.risk.append(
                    f"短期债务压力大：截至{debt_date}，发行人有息债务{round(interest_bearing_debt, 2)}亿元，其中短期债务{round(short_term_debt, 2)}亿元，占比{round(proportion * 100, 2)}%。")

    def risk_for_assets_liabilities(self, table):
        llm_prompt = [{
            "role": "user",
            "content": f"""/no_think
            你是金融文档分析专家，请从“资产负债表格”中提取关键信息，并以JSON格式输出：
            1. 提取最近一年的其他应收款，如果没有则不提取；
            1. 提取最近一年的存货，如果没有则不提取；
            2. 提取最近一年的流动资产，如果没有则不提取；
            3. 提取最近一年的资产总计，如果没有则不提取；
            4. 根据表格标题和表头最新的时间序列，提取报告截至日期，格式如'X年X月末'或'X年末'，如果没有则不提取；
            5. 结果以JSON格式输出；
            【资产负债表格】
            {table['data']}
            【表格标题】
            {table.get('preceding_text', "")}
            【结果输出】
            {{"其他应收款":0.0,"存货":0.0,"流动资产":0.0,"资产总计":0.0,"报告截至日期":""}}
            """
        }]
        _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
        if "亿" not in table['unit']:
            _result_json = self.unit_conversion(_result_json, table['unit'])
        assets_qita = float(_result_json["其他应收款"]) if _result_json["其他应收款"] else _result_json["其他应收款"]
        assets_cunhuo = float(_result_json["存货"]) if _result_json["存货"] else _result_json["存货"]
        assets_flow = float(_result_json["流动资产"]) if _result_json["流动资产"] else _result_json["流动资产"]
        assets_total = float(_result_json["资产总计"]) if _result_json["资产总计"] else _result_json["资产总计"]
        assets_time = _result_json["报告截至日期"] or "报告期末"
        if (assets_qita or assets_cunhuo) and assets_total:
            proportion = sum([assets_qita, assets_cunhuo]) / assets_total
            if proportion > 0.5:
                sub_str = ""
                risk_sub_str = ""
                if assets_qita and assets_cunhuo:
                    sub_str = f"其中其他应收款{round(assets_qita, 2)}亿元，存货{round(assets_cunhuo, 2)}亿元，"
                    risk_sub_str = "其他应收款和存货"
                elif assets_qita:
                    sub_str = f"其中其他应收款{round(assets_qita, 2)}亿元，"
                    risk_sub_str = "其他应收款"
                elif assets_cunhuo:
                    sub_str = f"其中存货{round(assets_cunhuo, 2)}亿元，"
                    risk_sub_str = "存货"
                self.risk.append(
                    f"资产流动性差：截至{assets_time}，发行人总资产{round(assets_total, 2)}亿元，{sub_str}合计占总资产比例{round(proportion * 100, 2)}%。")
                self.risk_evaluate.append(f"{risk_sub_str}合计占总资产比例超{round(proportion * 10 + 1) * 10}%")
        if assets_flow and assets_total:
            proportion_flow = float(assets_flow) / float(assets_total)
            if proportion_flow < 0.4:
                self.risk.append(
                    f"流动资产占比低：截至{assets_time}，发行人总资产{round(assets_total, 2)}亿元，其中流动资产{round(assets_flow, 2)}亿元，占比{round(proportion_flow * 100, 2)}%。")
                self.risk_evaluate.append(f"流动资产占比不足{round(proportion_flow * 10 - 1) * 10}%")
        return assets_total

    def risk_for_operating_revenue(self, table):
        llm_prompt = [{
            "role": "user",
            "content": f"""/no_think
            你是金融文档分析专家，请从“营业收入表格”中提取关键信息，并以JSON格式输出：
            1. 提取最近两年的年度总收入，以列表格式输出；
            2. 提取最近一期(非年度数据)总收入，如果没有则不提取；
            3. 提取最近两年的年度贸易收入，对应表述可能为贸易收入或者商品销售等，以列表格式输出，如果没有则不提取；
            4. 提取最近一期(非年度数据)贸易收入，对应表述可能为贸易收入或者商品销售等，如果没有则不提取；
            5. 结果以JSON格式输出；
            【营业收入表格】
            {table['data']}
            【结果输出】
            {{"最近两年的年度总收入":[], "最近一期总收入":0.0, "最近两年的年度贸易收入":[], "最近一期贸易收入":0.0}}
            """
        }]
        _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
        if "亿" not in table['unit']:
            _result_json = self.unit_conversion(_result_json, table['unit'])

        amounts = list(map(float, _result_json["最近两年的年度总收入"])) if _result_json["最近两年的年度总收入"] else _result_json["最近两年的年度总收入"]
        trade_amount = list(map(float, _result_json["最近两年的年度贸易收入"])) if _result_json["最近两年的年度贸易收入"] else _result_json["最近两年的年度贸易收入"]
        amounts_recently = float(_result_json["最近一期总收入"]) if _result_json["最近一期总收入"] else _result_json["最近一期总收入"]
        trade_amount_recently = float(_result_json["最近一期贸易收入"]) if _result_json["最近一期贸易收入"] else _result_json["最近一期贸易收入"]


        if len(amounts) == 2 and float(amounts[0]) / float(amounts[1]) < 0.9:
            amounts_ratio = round((1 - (amounts[0] / amounts[1])) * 100, 2)
            self.risk_evaluate.append("年度营业收入同比下降超10%")
            if amounts_recently:
                self.risk.append(
                    f"年度营业收入下降：近两年及一期，发行人分别实现营业收入{round(amounts[1], 2)}亿元、{round(amounts[0], 2)}亿元、{round(amounts_recently, 2)}亿元,年度营业收入同比下降{amounts_ratio}%。")
            else:
                self.risk.append(
                    f"年度营业收入下降：近两年，发行人分别实现营业收入{round(amounts[1], 2)}亿元、{round(amounts[0], 2)}亿元,年度营业收入同比下降{amounts_ratio}%。")
        if trade_amount and trade_amount[0] / amounts[0] > 0.3:
            trade_ratio = round(trade_amount[0] / amounts[0] * 100, 2)
            self.risk_evaluate.append("年度非贸易收入同比下降超10%")
            if trade_amount_recently:
                self.risk.append(
                    f"贸易收入占比高：近两年及一期，发行人分别实现营业收入{round(amounts[1], 2)}亿元、{round(amounts[0], 2)}亿元、{round(amounts_recently, 2)}亿元，其中贸易收入分别为{round(trade_amount_recently, 2)}亿元、{round(trade_amount[0], 2)}亿元、{round(trade_amount[1], 2)}亿元，占比{round(trade_amount_recently / amounts_recently * 100, 2)}%、{trade_ratio}%、{round(trade_amount[1] / amounts[1] * 100, 2)}%。")
            else:
                self.risk.append(
                    f"贸易收入占比高：近两年，发行人分别实现营业收入{round(amounts[1], 2)}亿元、{round(amounts[0], 2)}亿元，其中贸易收入分别为{round(trade_amount[0], 2)}亿元、{round(trade_amount[1], 2)}亿元，占比{trade_ratio}%、{round(trade_amount[1] / amounts[1] * 100, 2)}%。")

    def unit_conversion(self, input_str, unit):
        llm_prompt = [{
            "role": "user",
            "content": f"""
            你是一个专业的单位转换助手，你现在的任务是执行严格的数值单位转换，如果输入单位已经是亿(亿元)，则无需处理，直接返回原数据。如果单位不是亿(亿元)，请完全按照如下规则执行：
            1. 将输入JSON格式数据中包含的数字转换为成亿元为单位（如果需要）
            2. 数据可能包含分位符，转换前请先处理格式。
            3. 转换后必须进行结果校验，确认输入数据与结果数据之间的数量级关系正确。
            4. 结果以JSON格式输出，除转换的数字外，不做其他修改。
            【输入单位】
            {unit}
            【输入数据】
            {input_str}

            """
        }]
        _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
        return _result_json

    def search_by_keywords(self, keywords, lw_type=None):
        _data = []
        _txt_kws = []
        for _tp, _kw in keywords:
            if _tp == "tab" and self.table_datas.__contains__(_kw):
                _data.append(self.table_datas[_kw])
            elif _tp == "xml" and self.table_oxmls.__contains__(_kw):
                _data.append(self.table_oxmls[_kw])
            elif _tp == "txt":
                _txt_kws.append(_kw)
        if len(_txt_kws) == 0:
            return _data
        for _title, _context in self.article_datas.items():
            if lw_type == "without_guarantor" and exists_keyword(_title, ["(增信|信用增进)"]):
                continue
            # 内容为（x）开头的会被误认为标题而被过滤，这里添加此类内容
            if exists_keyword(_title, _txt_kws) and not _context and _title not in _data:
                _data.append(_title)
            if exists_keyword(_title, _txt_kws) and len(_context) > 0:
                _data.extend(_context)
        return _data

    def _get_reply_cust(self, issuer, reply_cust_list):
        for reply_cust in reply_cust_list:
            if issuer == reply_cust.get("custName", ""):
                return reply_cust.get("approvalNo", "")

    def gen_data_chapter2(self):
        """二、存量合作情况及我行政策"""
        logger.debug("二、存量合作情况及我行政策")
        _chapter_data = []

        _chapter_data.append({"data_src": [""], "type": "text", "context": "1、是否符合准入条件"})
        _chapter_data.append({"data_src": [""], "type": "dynamic_compliance_table", "context": ""})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "2、对发行人的存量授信情况，有无余额。"})

        issuer = self.company_info.get("发行人") or self.request_data.get("custName", "")

        # --- 新增：通过 API 获取数据 ---
        api_res = self.fetcher.get_credit_approval(issuer)
        latest_approval = {}
        if api_res.get("code") == 1000 and api_res.get("data"):
            datas = api_res["data"].get("datas", [])
            if datas:
                latest_approval = datas[0]  # 获取最新的一条批复

        reply_info = latest_approval.get("text_", "").strip()
        reply_date = latest_approval.get("reply_date", "").strip()
        # 映射 reply_type: 如果业务信息包含“集团”则设为 02，否则设为 01
        busi_info = latest_approval.get("busi_info_", "")
        reply_type = "02" if "集团" in busi_info else "01"
        # 映射 reply_cust: 使用 API 返回的批复号 approval_no 作为显示名称
        reply_cust = latest_approval.get("approval_no", "").strip()

        if reply_date:
            try:
                reply_date = datetime.strptime(reply_date, '%Y-%m-%d %H:%M:%S').strftime('%Y年%m月%d日')
            except:
                pass

        if reply_type == "02":
            llm_prompt = [{
                "role": "user",
                "content": f"""/no_think
                    你是金融文档分析专家，请从“银行授信批复信息”中提取关键信息，并以JSON格式输出：
                    1. 提取授信意见，从处理意见的内容开始提取，一直提取到集团授信详情前结束（如 具体授信方案如下：）；
                    2. 提取集团授信详情，以表格列表（二维数组）的格式输出（一般包含表头，数据项，合计三部分）；
                    3. 结果以JSON格式输出；
                    【银行授信批复信息】
                    {reply_info}
                    【城投公司名称】
                    {issuer}
                    【结果输出】
                    {{"授信意见":"", "授信详情": []}}
                    """
            }]
            _result_json = json.loads(self.private_apply_instance.generate(llm_prompt))
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"{reply_date}，我行" + _result_json.get("授信意见").strip()})
            _chapter_data.append({"data_src": [""], "type": "ltab", "context": _result_json.get("授信详情")})
        elif reply_type == "03" or reply_type == "01":
            if not reply_info:  # 容错处理：如果API没查到数据
                _chapter_data.append({"data_src": [""], "type": "text", "context": "新增客户"})
            elif issuer not in reply_info:
                _chapter_data.append(
                    {"data_src": [""], "type": "text", "context": f"【批复内容不包含客户（{issuer}）信息】"})
            elif reply_cust:
                _result = self.private_apply_instance.private_credit_approval(issuer, str(reply_info)[:30000],
                                                                              reply_date)
                _result = _result.replace("我行", f"{reply_cust}批复", 1)
                _chapter_data.append({"data_src": [""], "type": "text", "context": _result})
            else:
                _result = self.private_apply_instance.private_credit_approval(issuer, str(reply_info)[:30000],
                                                                              reply_date)
                _chapter_data.append({"data_src": [""], "type": "text", "context": _result})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "新增客户"})

        _param = self.search_by_keywords([("txt", "发行平台级别")])
        if len(_param) == 0:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "发行平台级别【无】"})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": _param})
        return _chapter_data

    def gen_data_chapter3(self):
        """三、区域经济情况"""
        logger.debug("三、区域经济情况")
        _chapter_data = []
        area_rank_list = self.search_by_keywords([('tab', '区域经济及债务')])
        area_rank_dict = area_rank_list[0] if len(area_rank_list) > 0 else {}
        if area_rank_dict == {}:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "区域经济情况【无】"})
            return _chapter_data
        _date = area_rank_dict['date']
        _unit = area_rank_dict['unit']
        area_rank_tab = [
            ['地区', 'GDP', '排名', '一般公共预算收入', '排名', '政府性基金收入', '排名', '宽口径债务率', '排名',
             '财政自给率', '排名']]
        if area_rank_dict['data']:
            area_rank_tab = area_rank_dict['data']
        # 文本信息
        _chapter_data.append({"data_src": [""], "type": "text",
                              "context": f"截至{custom_date(_date)}，区域经济及债务情况如下：{' ' * 6}单位：{_unit}元"})
        _chapter_data.append({"data_src": [""], "type": "ltab", "context": area_rank_tab})
        return _chapter_data

    def gen_data_chapter4(self):
        """四、拟投债券要素
        基本发行条款，授信申请方案，释义
        """
        logger.debug("四、拟投债券要素")
        _chapter_data = []
        _keywords = [('txt', '(实际控制人|释义)'), ('txt', '(解决机制|争议解决|违约责任|争议事项|违约事件)'),
                     ('txt', '募集资金.*?(用途|使用|运用计划)'), ('txt', '(主要|基本)*发行(条款|概况)'),
                     ('txt', '募集信息'), ('tab', '发行主体评级')]
        # 授信期限优先从调查报告中取
        _keywords_credit = [('txt', '授信申请方案')]
        _param_list = self.search_by_keywords(_keywords, lw_type="without_guarantor")
        _param_list_credit = self.search_by_keywords(_keywords_credit, lw_type="without_guarantor")
        _result = self.private_apply_instance.private_para4(
            str(_param_list)[:30000] + f"/n【授信申请方案】/n{_param_list_credit}")
        # 风险评价意见模块使用
        self.element = _result
        _chapter_data.append({"data_src": [""], "type": "text", "context": _result})
        return _chapter_data

    def gen_data_chapter5(self):
        """五、发行人概况"""
        logger.debug("五、发行人概况")
        _chapter_data = []

        # 1.发行人基本概况
        _keywords1 = [('txt', '发行人(基本)?(概况|情况|信息)'), ('txt', '(实际控制人|释义)'), ('txt', '授信申请方案'),
                      ('tab', '发行主体评级')]
        _param_list1 = self.search_by_keywords(_keywords1)
        if len(_param_list1) > 0:
            _result1 = self.private_apply_instance.private_faxing_condition(str(_param_list1)[:30000])
            _chapter_data.append({"data_src": [""], "type": "text", "context": _result1})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "发行人基本概况【无】"})

        # 1.工商信息
        _chapter_data.append({"data_src": [""], "type": "text", "context": "1、工商信息"})
        try:
            basic_info_data = self.comprehensive_generator.data.get("公司基本信息", {})
            if basic_info_data:
                # 构造工商信息表格数据
                basic_info_fields = [
                    ["注册名称", basic_info_data.get("COMPNAME", "")],
                    ["法定代表人", basic_info_data.get("LEGREP", "")],
                    ["注册资本",
                     f"{basic_info_data.get('REGCAPITAL', '')} 万元" if basic_info_data.get('REGCAPITAL') else ""],
                    ["设立日期", basic_info_data.get("FOUNDDATE", "")],
                    ["统一社会信用代码", basic_info_data.get("unified_credit_code", "")],
                    ["公司住所", basic_info_data.get("REGADDR", "")],
                    ["办公地址", basic_info_data.get("OFFICEADDR", "")],
                    ["邮政编码", basic_info_data.get("OFFICEZIPCODE", "")],
                    ["电话", basic_info_data.get("COMPTEL", "")],
                    ["传真", basic_info_data.get("COMPFAX", "")],
                    ["经营范围", basic_info_data.get("BIZSCOPE", "")]
                ]
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": basic_info_fields})
            else:
                _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无公司基本信息】"})
        except Exception as e:
            logger.exception(f"获取工商信息失败: {e}")
            _chapter_data.append({"data_src": [""], "type": "text", "context": "工商信息获取失败"})

        # 2.股东信息
        _chapter_data.append({"data_src": [""], "type": "text", "context": "2、股权结构"})
        try:
            # 检索股权结构相关内容
            _keywords_equity = [('txt', '(股权结构|股东|持股|实际控制人|控股股东|出资|资本构成|股权穿透)')]
            _param_list_equity = self.search_by_keywords(_keywords_equity)

            # 调用private_apply中的方法生成描述
            equity_description = self.private_apply_instance.get_equity_structure_description(_param_list_equity)

            if equity_description:
                _chapter_data.append({"data_src": [""], "type": "text", "context": equity_description})
        except Exception as e:
            logger.exception(f"获取股权结构描述失败: {e}")
        try:
            shareholders = self.comprehensive_generator.data.get("股东信息", [])
            if shareholders:
                # 构造股东表格,包含表头
                shareholder_table = [["股东名称", "股东内码", "持股数量", "持股比例"]]
                for item in shareholders:
                    shareholder_table.append([
                        item.get("SHHOLDERNAME", ""),
                        item.get("SHHOLDERSECODE", ""),
                        item.get("HOLDERAMT", ""),
                        f"{item.get('HOLDERRTO', '')}%" if item.get('HOLDERRTO') else ""
                    ])
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": shareholder_table})
            else:
                _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无股东信息】"})
        except Exception as e:
            logger.exception(f"获取股东信息失败: {e}")
            _chapter_data.append({"data_src": [""], "type": "text", "context": "股东信息获取失败"})

        # 4.主体评级列表(新增)
        _chapter_data.append({"data_src": [""], "type": "text", "context": "3、主体评级"})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "发行人主体评级信息如下："})
        try:
            ratings = self.comprehensive_generator.data.get("主体评级列表", [])
            if ratings:
                # 构造评级表格,包含表头
                rating_table = [["评级公司", "级别日期", "级别", "展望", "披露日期", "有效截止日", "级别对象"]]
                for item in ratings:
                    rating_table.append([
                        item.get("RATECOMNAME", ""),
                        item.get("PUBLISHDATE", ""),
                        item.get("CREDITRATE", ""),
                        item.get("EXPTRATING_value", ""),
                        item.get("DECLAREDATE", ""),
                        item.get("CREDITRATEENDDATE", ""),
                        item.get("COMTYPE_value", "")
                    ])
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": rating_table})
            else:
                _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无评级信息】"})
        except Exception as e:
            logger.exception(f"获取评级信息失败: {e}")
            _chapter_data.append({"data_src": [""], "type": "text", "context": "评级信息获取失败"})

        # 5.区域城投债平台排名
        _chapter_data.append({"data_src": [""], "type": "text", "context": "4、区域排名"})
        _keywords2 = [('tab', '发债平台情况')]
        # _param_list2 = self.search_by_keywords(_keywords2)
        # if _param_list2 and isinstance(_param_list2[0], dict):
        #     _result2 = self.private_apply_instance.private_fazhai_rank(str(_param_list2))
        #     _chapter_data.append({"data_src": [""], "type": "text", "context": _result2})
        # else:
        #     _result2 = "【没有检索到区域城投债平台排名数据】"
        #     _chapter_data.append({"data_src": [""], "type": "text", "context": _result2})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "发行人区域排名情况如下："})
        try:
            platform_scores = self.comprehensive_generator.data.get("区域排名信息", [])
            if platform_scores:
                platform_table = [["公司名称", "区域", "城投评分", "省内排名", "主体评级", "债券余额"]]
                for item in platform_scores:
                    platform_table.append([
                        item.get("itname", ""),
                        item.get("regionname", ""),
                        item.get("score_all", ""),
                        item.get("rank", ""),
                        item.get("credit_rate", ""),
                        item.get("bondbalance", "")
                    ])
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": platform_table})
        except Exception as e:
            logger.debug(f"从structured获取平台排名失败: {e}")

        # 6.存量债券
        _chapter_data.append({"data_src": [""], "type": "text", "context": "5、存量债券"})
        _keywords3 = [('tab', '存量债券'), ('tab', '存量债券统计')]
        _param_list3 = self.search_by_keywords(_keywords3)
        if _param_list3 == [] or _param_list3 == [{}]:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "存量债券明细【无】"})
        else:
            _result3 = self.private_apply_instance.private_zhaiquan_yue_v1(str(_param_list3))
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": (
                                                 '' if _result3 is None else _result3) + f"存量债券明细如下:{' ' * 6}单位:{_param_list3[0]['unit']}元、%"})
            _chapter_data.append({"data_src": [""], "type": "ltab", "context": _param_list3[0]['data']})


        # 6.信用利差
        _chapter_data.append({"data_src": [""], "type": "text", "context": "6、信用利差"})
        try:
            spreads = self.comprehensive_generator.data.get("利差信息", [])
            if spreads:
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"{self.company_name} 近七天平均信用利差如下:"})
                spread_table = [["交易日期", "当前利差(BP)"]]
                for item in spreads:
                    spread_table.append([
                        item.get("tradedate", ""),
                        item.get("spread", "")
                    ])
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": spread_table})
            else:
                _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无利差信息】"})
        except Exception as e:
            logger.exception(f"获取利差信息失败: {e}")
            _chapter_data.append({"data_src": [""], "type": "text", "context": "利差信息获取失败"})

        # 7.注册批复全景
        _chapter_data.append({"data_src": [""], "type": "text", "context": "7、批文情况"})
        try:
            registrations = self.comprehensive_generator.data.get("注册批复全景", [])
            if registrations:
                reg_table = [["项目名称", "批复场所", "品种", "注册中额度", "最新注册时间", "最新注册状态"]]
                for item in registrations:
                    reg_table.append([
                        item.get("project_name", ""),
                        item.get("approval_location", ""),
                        item.get("project_type_name", ""),
                        item.get("register_amount", ""),
                        item.get("process_date", ""),
                        item.get("process_type_name", "")
                    ])
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": reg_table})
            else:
                _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无注册批复信息】"})
        except Exception as e:
            logger.exception(f"获取注册批复信息失败: {e}")
            _chapter_data.append({"data_src": [""], "type": "text", "context": "注册批复信息获取失败"})

        # 9.主营业务总结 暂不需要
        _param_list5 = self.search_by_keywords([('tab', '营业收入')])
        if len(_param_list5) > 0 and _param_list5[0]:
            _result5 = self.private_apply_instance.private_5para_yinye_shouru(str(_param_list5))
            param_dict5 = _param_list5[0] if len(_param_list5) > 0 else {}
            _unit5 = param_dict5['unit']

        # 10.主营业务表格
        _chapter_data.append({"data_src": [""], "type": "text", "context": "8、营业收入"})
        _keywords6 = [('xml', '营业收入')]
        _param_list6 = self.search_by_keywords(_keywords6)
        if len(_param_list6) > 0 and _param_list6[0]:
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"报告期内,发行人营业收入情况如下:{' ' * 6}单位:{_unit5}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list6[0]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无营业收入数据】"})

        return _chapter_data

    # def gen_data_chapter6(self):
    #     """六、财务情况"""
    #     logger.debug("六、财务情况")
    #     _chapter_data = []
    #
    #     # 1.资产概况分析
    #     _keywords1 = [('tab', '资产负债表')]
    #     _param_list1 = self.search_by_keywords(_keywords1)
    #     _result1 = "资产详情【无】"
    #     if not _param_list1:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result1})
    #     else:
    #         _result1 = self.private_apply_instance.private_6para_first(str(_param_list1)[:30000])
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result1})
    #     # 2.资产科目分析
    #     # 与资产概况分析合并
    #
    #     # 3.应收账款
    #     _keywords3 = [('xml', '应收账款')]
    #     _param_list3 = self.search_by_keywords(_keywords3)
    #     if not _param_list3:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "应收账款表格详情【无】"})
    #     elif _param_list3 is not None and len(_param_list3) > 0 and _param_list3[0].__contains__("oxml") and _param_list3[0]['oxml'] is not None:
    #         _chapter_data.append({"data_src": [""], "type": "text",
    #                               "context": f"截至{custom_date(_param_list3[0]['date'])}，应收账款主要明细如下：{' ' * 6}单位：{_param_list3[0]['unit']}"})
    #         _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list3[0]['oxml']})
    #
    #     # 5.其他应收款
    #     _keywords5 = [('xml', '其他应收款')]
    #     _param_list5 = self.search_by_keywords(_keywords5)
    #     if not _param_list5:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "其他应收款表格详情【无】"})
    #     elif _param_list5 is not None and len(_param_list5) > 0 and _param_list5[0].__contains__("oxml") and _param_list5[0]['oxml'] is not None:
    #         _chapter_data.append({"data_src": [""], "type": "text",
    #                               "context": f"截至{custom_date(_param_list5[0]['date'])}，其他应收款主要明细如下：{' ' * 6}单位：{_param_list5[0]['unit']}"})
    #         _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list5[0]['oxml']})
    #
    #     # 6.受限资产
    #     _keywords6 = [('xml', '受限资产')]
    #     _param_list6 = self.search_by_keywords(_keywords6)
    #     if not _param_list6:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "受限资产表格详情【无】"})
    #     elif _param_list6 is not None and len(_param_list6) > 0 and _param_list6[0].__contains__("oxml") and _param_list6[0]['oxml'] is not None:
    #         _chapter_data.append({"data_src": [""], "type": "text",
    #                               "context": f"截至{custom_date(_param_list6[0]['date'])}，受限资产情况如下：{' ' * 6}单位：{_param_list6[0]['unit']}"})
    #         _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list6[0]['oxml']})
    #
    #     # 7.负债分析 负债统计
    #     _keywords7 = [('tab', '资产负债表')]
    #     _param_list7 = self.search_by_keywords(_keywords7)
    #     _result7 = None
    #     if not _param_list7:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "负债详情【无】"})
    #     else:
    #         _result7 = self.private_apply_instance.private_6para_fuzhai(str(_param_list7)[:30000])
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result7})
    #
    #     # 8.有息负债
    #     _keywords8 = [('tab', '有息负债-募集')]
    #     _param_list8 = self.search_by_keywords(_keywords8)
    #     _result8 = None
    #     if not _param_list8:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "负债详情【无】"})
    #     else:
    #         # 去掉干扰上下文
    #         try:
    #             table_data = _param_list8[0].get("data")
    #             if table_data[0][0] == "项目" and "合计" in table_data[-1][0]:
    #                 change_flag = False
    #                 res_table = []
    #                 for index, row in enumerate(table_data):
    #                     if row[0] == "项目":
    #                         res_table.append(row)
    #                     elif row[0] == "银行贷款":
    #                         change_flag = True
    #                         break
    #                     else:
    #                         break
    #                 if change_flag:
    #                     res_table.append(table_data[-1])
    #                     _param_list8[0]["data"] = res_table
    #         except Exception as e:
    #             logger.exception(e)
    #         _result8 = self.private_apply_instance.private_6para_youxifuzhai(str(_param_list8)[:30000],
    #                                                                          self.fuzhai_total)
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result8})
    #
    #     # 9.融资结构 = 有息负债
    #     _keywords9 = [('xml', '有息负债')]
    #     _param_list9 = self.search_by_keywords(_keywords9)
    #     # logger.debug("资产负债结果：", _result6)
    #     if _param_list9 is not None and len(_param_list9) > 0 and _param_list9[0].__contains__("oxml") and _param_list9[0]['oxml'] is not None:
    #         _chapter_data.append({"data_src": [""], "type": "text",
    #                               "context": f"截至{custom_date(_param_list9[0]['date'])}，融资结构情况如下：{' ' * 6}单位：{_param_list9[0]['unit']}"})
    #         _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list9[0]['oxml']})
    #
    #     # 10.授信机构
    #     _keywords10 = [('tab', '授信情况'), ('xml', '授信情况')]
    #     _param_list10 = self.search_by_keywords(_keywords10)
    #     _result10 = None
    #     if not _param_list10 or len(_param_list10) <= 1 or not _param_list10[1].__contains__('oxml') or not _param_list10[1]['oxml'] is not None:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "授信详情【无】"})
    #     else:
    #         _result10 = self.private_apply_instance.private_6para_shouxinedu(str(_param_list10[0])[:30000])
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": "" if _result10 is None else _result10})
    #
    #     # 11.现金流
    #     _keywords11 = [('tab', '现金流')]
    #     _param_list11 = self.search_by_keywords(_keywords11)
    #     _result11 = "现金流详情【无】"
    #     if not _param_list11:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result11})
    #     else:
    #         _result11 = self.private_apply_instance.private_6para_xianjinliu(str(_param_list11)[:30000])
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result11})
    #
    #     # 12.对外担保
    #     _keywords12 = [('txt', '对外担保')]
    #     _param_list12 = self.search_by_keywords(_keywords12, lw_type="without_guarantor")
    #     if not _param_list12:
    #         _param_list12 = self.search_by_keywords([('tab', '对外担保')], lw_type="without_guarantor")
    #     _result12 = "对外担保详情【无】"
    #     if not _param_list12:
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result12})
    #     else:
    #         _result12 = self.private_apply_instance.private_6para_danbao(str(_param_list12)[:30000])
    #         _chapter_data.append({"data_src": [""], "type": "text", "context": _result12})
    #     return _chapter_data

    def gen_data_chapter6(self):
        """六、财务情况"""
        logger.debug("六、财务情况")
        _chapter_data = []

        # ========== 第（一）部分 财务报表 ==========
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": [{"bold": True, "content": "第（一）部分 财务报表"}]})

        # 1、资产负债表
        _chapter_data.append({"data_src": [""], "type": "text", "context": "1、资产负债表"})
        _keywords_zcfzb = [('xml', '资产负债表')]
        _param_list_zcfzb = self.search_by_keywords(_keywords_zcfzb)
        if _param_list_zcfzb and len(_param_list_zcfzb) > 0 and _param_list_zcfzb[0].get("oxml"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list_zcfzb[0]['date'])}，资产负债表如下：{' ' * 6}单位：{_param_list_zcfzb[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list_zcfzb[0]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无资产负债表数据】"})

        # 2、利润表
        _chapter_data.append({"data_src": [""], "type": "text", "context": "2、利润表"})
        _keywords_lrb = [('xml', '利润表')]
        _param_list_lrb = self.search_by_keywords(_keywords_lrb)
        if _param_list_lrb and len(_param_list_lrb) > 0 and _param_list_lrb[0].get("oxml"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list_lrb[0]['date'])}，利润表如下：{' ' * 6}单位：{_param_list_lrb[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list_lrb[0]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无利润表数据】"})

        # ========== 第（二）部分 重点科目分析 ==========
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": [{"bold": True, "content": "第（二）部分 重点科目分析"}]})

        # 1、资产情况
        _chapter_data.append({"data_src": [""], "type": "text", "context": "1、资产情况"})

        # 资产概况分析
        _keywords1 = [('tab', '资产负债表')]
        _param_list1 = self.search_by_keywords(_keywords1)
        if _param_list1:
            _result1 = self.private_apply_instance.private_6para_first(str(_param_list1)[:30000])
            _chapter_data.append({"data_src": [""], "type": "text", "context": _result1})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "资产详情【无】"})

        # 应收账款
        _chapter_data.append({"data_src": [""], "type": "text", "context": "应收账款："})
        _keywords3 = [('xml', '应收账款')]
        _param_list3 = self.search_by_keywords(_keywords3)
        if _param_list3 and len(_param_list3) > 0 and _param_list3[0].get("oxml"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list3[0]['date'])}，应收账款主要明细如下：{' ' * 6}单位：{_param_list3[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list3[0]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无应收账款数据】"})

        # 其他应收款
        _chapter_data.append({"data_src": [""], "type": "text", "context": "其他应收款："})
        _keywords5 = [('xml', '其他应收款'), ('tab', '其他应收款')]
        _param_list5 = self.search_by_keywords(_keywords5)
        if _param_list5 and len(_param_list5) > 0:
            # 优先使用 oxml 格式（保留完整表格格式）
            if _param_list5[0].get("oxml"):
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"截至{custom_date(_param_list5[0]['date'])}，其他应收款明细如下：{' ' * 6}单位：{_param_list5[0]['unit']}"})
                _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list5[0]['oxml']})
            # 如果没有 oxml，使用 data 格式（完整二维数组）
            elif _param_list5[0].get("data"):
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"截至{custom_date(_param_list5[0]['date'])}，其他应收款明细如下：{' ' * 6}单位：{_param_list5[0]['unit']}"})
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": _param_list5[0]['data']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无其他应收款数据】"})

        # 受限资产
        _chapter_data.append({"data_src": [""], "type": "text", "context": "受限资产："})
        _keywords6 = [('xml', '受限资产')]
        _param_list6 = self.search_by_keywords(_keywords6)
        if _param_list6 and len(_param_list6) > 0 and _param_list6[0].get("oxml"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list6[0]['date'])}，受限资产情况如下：{' ' * 6}单位：{_param_list6[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list6[0]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无受限资产数据】"})

        # 2、负债情况
        _chapter_data.append({"data_src": [""], "type": "text", "context": "2、负债情况"})

        # 负债分析
        _keywords7 = [('tab', '资产负债表')]
        _param_list7 = self.search_by_keywords(_keywords7)
        if _param_list7:
            _result7 = self.private_apply_instance.private_6para_fuzhai(str(_param_list7)[:30000])
            _chapter_data.append({"data_src": [""], "type": "text", "context": _result7})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "负债详情【无】"})

        # 有息债务
        _chapter_data.append({"data_src": [""], "type": "text", "context": "有息债务："})
        _keywords8 = [('tab', '有息负债-募集')]
        _param_list8 = self.search_by_keywords(_keywords8)
        if _param_list8:
            # 去掉干扰上下文
            try:
                table_data = _param_list8[0].get("data")
                if table_data and table_data[0][0] == "项目" and "合计" in table_data[-1][0]:
                    change_flag = False
                    res_table = []
                    for index, row in enumerate(table_data):
                        if row[0] == "项目":
                            res_table.append(row)
                        elif row[0] == "银行贷款":
                            change_flag = True
                            break
                        else:
                            break
                    if change_flag:
                        res_table.append(table_data[-1])
                        _param_list8[0]["data"] = res_table
            except Exception as e:
                logger.exception(e)
            _result8 = self.private_apply_instance.private_6para_youxifuzhai(str(_param_list8)[:30000],
                                                                             self.fuzhai_total)
            _chapter_data.append({"data_src": [""], "type": "text", "context": _result8})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无有息债务数据】"})

        # 融资结构(有息负债表格)
        _keywords9 = [('xml', '有息负债')]
        _param_list9 = self.search_by_keywords(_keywords9)
        if _param_list9 and len(_param_list9) > 0 and _param_list9[0].get("oxml"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list9[0]['date'])}，融资结构情况如下：{' ' * 6}单位：{_param_list9[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list9[0]['oxml']})

        # 非标融资
        _chapter_data.append({"data_src": [""], "type": "text", "context": "非标融资："})
        _keywords_feibiao = [('tab', '非标融资')]  # 只需要 tab 类型
        _param_list_feibiao = self.search_by_keywords(_keywords_feibiao)
        if _param_list_feibiao and len(_param_list_feibiao) > 0 and _param_list_feibiao[0].get("data"):
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"截至{custom_date(_param_list_feibiao[0]['date'])}，非标融资情况如下：{' ' * 6}单位：{_param_list_feibiao[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "ltab", "context": _param_list_feibiao[0]['data']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无非标融资数据】"})

        # 3、发行人金融机构授信情况
        _chapter_data.append({"data_src": [""], "type": "text", "context": "3、发行人金融机构授信情况"})
        _keywords10 = [('tab', '授信情况'), ('xml', '授信情况')]
        _param_list10 = self.search_by_keywords(_keywords10)
        if _param_list10 and len(_param_list10) > 0:
            # 先输出文字分析
            _result10 = self.private_apply_instance.private_6para_shouxinedu(str(_param_list10[0])[:30000])
            if _result10:
                _chapter_data.append({"data_src": [""], "type": "text", "context": _result10})

            # 再输出表格
            if len(_param_list10) > 1 and _param_list10[1].get('oxml'):
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"截至{custom_date(_param_list10[1]['date'])}，授信情况如下：{' ' * 6}单位：{_param_list10[1]['unit']}"})
                _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list10[1]['oxml']})
        else:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "【暂无授信数据】"})

        # 11.现金流
        _chapter_data.append({"data_src": [""], "type": "text", "context": "4、现金流情况"})
        _keywords11 = [('xml', '现金流')]
        _param_list11 = self.search_by_keywords(_keywords11)
        _result11 = "现金流详情【无】"
        if not _param_list11:
            _chapter_data.append({"data_src": [""], "type": "text", "context": _result11})
        else:
            _result11 = self.private_apply_instance.private_6para_xianjinliu(str(_param_list11)[:30000])
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": f"发行人现金流情况如下：{' ' * 6}单位：{_param_list11[0]['unit']}"})
            _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list11[0]['oxml']})

        _chapter_data.append({"data_src": [""], "type": "text", "context": "5、对外担保情况"})

        # 优先获取表格数据（xml格式保留完整表格样式，tab格式为二维数组）
        _keywords12_xml = [('xml', '对外担保')]
        _keywords12_tab = [('tab', '对外担保')]
        _param_list12_xml = self.search_by_keywords(_keywords12_xml, lw_type="without_guarantor")
        _param_list12_tab = self.search_by_keywords(_keywords12_tab, lw_type="without_guarantor")

        # 同时获取文本数据用于LLM总结
        _keywords12_txt = [('txt', '对外担保')]
        _param_list12_txt = self.search_by_keywords(_keywords12_txt, lw_type="without_guarantor")

        # 用于LLM总结的数据（优先使用文本，其次使用表格）
        _param_for_summary = _param_list12_txt if _param_list12_txt else _param_list12_tab

        if not _param_for_summary and not _param_list12_xml and not _param_list12_tab:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "对外担保详情【无】"})
        else:
            # 先输出LLM总结
            if _param_for_summary:
                _result12 = self.private_apply_instance.private_6para_danbao(str(_param_for_summary)[:30000])
                _chapter_data.append({"data_src": [""], "type": "text", "context": _result12})

            # 再输出完整的对外担保表格
            if _param_list12_xml and len(_param_list12_xml) > 0 and _param_list12_xml[0].get("oxml"):
                # 使用xml格式输出（保留原始表格样式）
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"截至{custom_date(_param_list12_xml[0].get('date', ''))}，对外担保明细如下：{' ' * 6}单位：{_param_list12_xml[0].get('unit', '万元')}"})
                _chapter_data.append({"data_src": [""], "type": "xtab", "context": _param_list12_xml[0]['oxml']})
            elif _param_list12_tab and len(_param_list12_tab) > 0 and _param_list12_tab[0].get("data"):
                # 使用tab格式输出（二维数组表格）
                _chapter_data.append({"data_src": [""], "type": "text",
                                      "context": f"截至{custom_date(_param_list12_tab[0].get('date', ''))}，对外担保明细如下：{' ' * 6}单位：{_param_list12_tab[0].get('unit', '万元')}"})
                _chapter_data.append({"data_src": [""], "type": "ltab", "context": _param_list12_tab[0]['data']})

        return _chapter_data


    def gen_data_chapter7(self):
        """七、保证人"""
        _chapter_data = []
        # 根据募集说明书开头文本判断，如果无增信，直接返回无保证人
        if self.company_info.get("是否存在增信", "否") == "否":
            _chapter_data.append({"data_src": [""], "type": "text", "context": '无'})
            return _chapter_data

        # 根据保证人文本判断
        _keywords1 = [('txt', '(增信|信用增进)'), ('tab', '保证人基本情况')]
        _param_list1 = self.search_by_keywords(_keywords1)
        if len(_param_list1) == 0:
            _chapter_data.append({"data_src": [""], "type": "text", "context": '无'})
            return _chapter_data

        _result1 = self.private_apply_instance.private_6para_baozhengren(str(_param_list1)[:30000])
        context = '无' if _result1 is None else _result1
        _chapter_data.append({"data_src": [""], "type": "text", "context": context})
        return _chapter_data

    def gen_data_chapter8(self):
        """八、风险点"""
        _chapter_data = []
        if not self.risk:
            _chapter_data.append({"data_src": [""], "type": "text", "context": "无"})
        for index, evaluate in enumerate(self.risk):
            _chapter_data.append({"data_src": [""], "type": "text", "context": f"{index + 1}、{evaluate}"})
        return _chapter_data

    def gen_data_chapter9(self):
        """九、风险评价人意见"""
        name = re.search(r'债券名称[：:](.*)\n', self.element).group(1).strip().strip("。").strip("，")
        deadline = re.search(r'发行期限[：:](.*)\n', self.element).group(1).strip().strip("。").strip("，")
        # deadline = re.search(r'发行期限[:：](.*)\n', self.element).group(1).strip()
        rate = re.search(r'票面利率[：:](.*)\n', self.element).group(1).strip().strip("。").strip("，")
        grade = re.search(r'外部评级[：:](.*)\n', self.element).group(1).strip().strip("。").strip("，")
        guarantee_method = re.search(r'担保方式[：:](.*)\n', self.element).group(1).strip().strip("。").strip("，")
        involve = "发行人"
        if self.company_info.get("担保人", None):
            involve = "发行人和保证人"

        _chapter_data = []
        _chapter_data.append({"data_src": [""], "type": "text", "context": f'拟申请给予{self.company_info.get("发行人", "【发行人解析错误】")}授信敞口额度【xxx】，授信额度有效期一年，额度可循环使用。'})
        _chapter_data.append({"data_src": [""], "type": "text", "context": [{"bold": True, "content": "授信品种："}, {
            "content": f"债券投资，用于投资“{name}”。额度可调剂用于相同发行人发行的期限不超过5年的其他债券，单笔投资金额不超过实际发行金额的20%。"}]})
        _chapter_data.append({"data_src": [""], "type": "text", "context": [{"bold": True, "content": "投资期限："}, {
            "content": f"{deadline}，我行投资不超过2年。"}]})
        # _chapter_data.append({"data_src": [""], "type": "text", "context": [{"bold":True, "content":"发行日期："}, {"content":f"2025年7月24日"}]})
        _chapter_data.append({"data_src": [""], "type": "text", "context": [{"bold": True, "content": "票面利率："}, {
            "content": f"{rate}，具体定价须经我行财务部门审核同意"}]})
        if self.company_info.get("担保人", None) and guarantee_method:
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": [{"bold": True, "content": "担保方式："}, {"content": guarantee_method}]})
        else:
            _chapter_data.append({"data_src": [""], "type": "text",
                                  "context": [{"bold": True, "content": "担保方式："}, {"content": f"信用"}]})
        # _chapter_data.append(
        #     {"data_src": [""], "type": "text", "context": [{"bold": True, "content": f"（一）放款前提条件："}]})
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'1、金融市场部认真核实发行文件，确保债券要素与申报材料信息一致，{involve}主体评级无下调。'})
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'2、金融市场部放款前再次通过公开渠道查询交易对手外部公开信息，确保关联公司无违约等重大负面信息。'})
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'3、{self.company_info.get("发行人", "【发行人解析错误】")}授信额度【xxx】，如调剂用于相同发行人发行的期限不超过{deadline}的其他债券，用信需经金融市场部风险总监签批同意。'})
        # _chapter_data.append(
        #     {"data_src": [""], "type": "text", "context": [{"bold": True, "content": f"（二）投后管理要求："}]})
        # if self.company_info.get("担保人", None):
        #     _chapter_data.append({"data_src": [""], "type": "text",
        #                           "context": f'1、金融市场部按季关注发行人公开披露的报表情况，并关注外部给予的评级。若发生以下情况之一，则禁止投资，且将持有债券择机卖出：（1）发行人和保证人主体评级下降或展望负面；（2）发生其他对发行人偿债能力有重大影响的事件。'})
        # else:
        #     _chapter_data.append({"data_src": [""], "type": "text",
        #                           "context": f'1、金融市场部按季关注发行人公开披露的报表情况，并关注外部给予的评级。若发生以下情况之一，则禁止投资，且将持有债券择机卖出：（1）发行人主体评级下降或展望负面；（2）发生其他对发行人偿债能力有重大影响的事件。'})
        #
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'2、金融市场部按照《江苏苏商银行金融市场业务投后管理办法》相关要求进行投后管理。按季收集定期报告、临时报告、债项及主体评级报告等，并对发行人信用状况进行评估，如有不利情况，须及时预警并反馈风险管理部，确保我行资产安全。'})
        #
        # if self.risk_evaluate:
        #     _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'3、金融市场部关注发行人经营及财务状况、偿债能力、经营现金流情况、债券估值，如{"、".join(self.risk_evaluate)}或出现违约，须及时预警，且不适用放款前提条件3涉及的调剂规则。'})
        # else:
        #     _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'3、金融市场部关注发行人经营及财务状况、偿债能力、经营现金流情况、债券估值，如出现违约，须及时预警，且不适用放款前提条件3涉及的调剂规则。'})
        #
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'4、金融市场部持有期内可视该债券市场估值变化，在符合我行规程的前提下择机转出，我行持有期最长不超过2年。'})
        # _chapter_data.append({"data_src": [""], "type": "text", "context": f'未尽事宜按监管规定及我行管理制度执行。'})
        # _chapter_data.append({"data_src": [""], "type": "text",
        #                       "context": f'本风险评价意见书是基于金融市场部尽职调查所提供的资料以及问题反馈材料得出的结论，作为主审本人对结论的真实性和完整性负责。'})
        return _chapter_data

    def gen_data_chapter10(self):
        """十一、授信调查声明"""
        _chapter_data = []

        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "一、本授信业务的前期调查、核实工作由本人亲自完成。"})
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "二、贷前调查报告所反映的信息均通过采用了实地调查方式获取。"})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "三、我们都不是授信借款申请人的关系人。"})
        _chapter_data.append({"data_src": [""], "type": "text",
                              "context": "四、我们确认授信调查报告所陈述事实和授信调查结论的真实性，并承担由于产生的一切后果。"})
        _chapter_data.append({"data_src": [""], "type": "text",
                              "context": "五、我们确认贷前调查报告所附资料的真实性，资料如为复印均已与原件核对无误，并承担由于资料失实引致的一切后果。"})

        # 添加空行
        _chapter_data.append({"data_src": [""], "type": "text", "context": "\n\n"})

        # 签名区域
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "                                  主办客户经理签名："})
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "                                              日期："})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "\n"})

        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "                                  协办客户经理签名："})
        _chapter_data.append(
            {"data_src": [""], "type": "text", "context": "                                              日期："})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "\n"})

        _chapter_data.append({"data_src": [""], "type": "text", "context": "                                  部门负责人签名："})
        _chapter_data.append({"data_src": [""], "type": "text", "context": "                                              日期："})

        return _chapter_data

    def gen_report_datas(self):
        """
        生成报告数据
        [{"title":"","paragraphs":[{"data_src":["路径"],"type":"table/text","context":""}]}]
        """
        report_datas = []
        # 数据准备
        self.prepare_data()
        # 一、合作方案
        report_datas.append({"title": "一、合作方案", "paragraphs": [{"data_src": [""], "type": "text", "context": "\n\n\n\n\n"}]})
        # 二、存量合作情况及我行政策
        report_datas.append({"title": "二、存量合作情况及我行政策", "paragraphs": self.gen_data_chapter2()})
        # 三、区域经济情况
        report_datas.append({"title": "三、区域经济情况", "paragraphs": self.gen_data_chapter3()})
        # 四、拟投债券要素
        report_datas.append({"title": "四、拟投债券要素", "paragraphs": self.gen_data_chapter4()})
        # 五、发行人概况
        report_datas.append({"title": "五、发行人概况", "paragraphs": self.gen_data_chapter5()})
        # 六、财务情况
        report_datas.append({"title": "六、财务情况", "paragraphs": self.gen_data_chapter6()})
        # 七、保证人
        report_datas.append({"title": "七、保证人", "paragraphs": self.gen_data_chapter7()})
        # 八、优势
        report_datas.append({"title": "八、优势", "paragraphs": [{"data_src": [""], "type": "text", "context": "\n\n\n\n\n"}]})
        # 九、风险点
        report_datas.append({"title": "九、风险点", "paragraphs": self.gen_data_chapter8()})
        # 十、发起部门意见
        report_datas.append({"title": "十、发起部门意见", "paragraphs": self.gen_data_chapter9()})
        # 十一、授信调查声明
        report_datas.append({"title": "                        授信调查声明", "paragraphs": self.gen_data_chapter10()})

        return report_datas

    def add_paragraph(self, target_doc, chapter_data):
        _title = chapter_data['title']
        add_one_level_header(target_doc, _title)
        logger.debug(_title)
        _paragraphs = chapter_data['paragraphs']

        for _para_item in _paragraphs:
            if "xtab" == _para_item['type']:
                create_table_by_oxml(target_doc, _para_item['context'])
            elif "ltab" == _para_item['type']:
                create_table(target_doc, _para_item['context'])
            elif "dynamic_compliance_table" == _para_item['type']:
                # 创建表格结构
                table = self.comprehensive_generator._create_compliance_table_structure(target_doc)
                # 填充数据
                self.comprehensive_generator._fill_compliance_data(table)
            else:
                if isinstance(_para_item['context'], str):
                    for context in _para_item['context'].split('\n'):
                        add_content(target_doc, context)
                else:
                    add_content(target_doc, _para_item['context'])

    def gen_report(self):
        try:
            """生成报表"""
            self.start_time = time.time()
            # 获取数据
            report_datas = self.gen_report_datas()

            target_doc = Document()

            add_cover_page(target_doc)

            # 添加内容
            # [{"title": "", "paragraphs": [{"data_src": ["路径"], "type": "table/text", "context": ""}]}]
            for chapter_data in report_datas:
                self.add_paragraph(target_doc, chapter_data)
            # 输出文档
            output_file = f"信用债调查报告-{self.company_name}.docx"
            c_output = os.path.join(self.root_data_dir, output_file)
            target_doc.save(c_output)
            logger.info(f"gen report: {output_file} cost time：{time.time() - self.start_time}")
            return c_output
        except Exception as e:
            logger.exception(e)
            raise Exception(f"报告生成错误。{str(e)}")


if __name__ == "__main__":
    # company_name = "德阳经开区发展（控股）集团有限公司"
    root_data_dir = "/home/datagov/INSTANCE/projects/report_xyzdc/files/test1"

    # report = UrbanReport(company_name, root_data_dir, output_dir)
    # report.gen_report()
    company_name = "都江堰兴市集团有限责任公司"
    report = UrbanReport(company_name, root_data_dir)
    report.gen_report()
