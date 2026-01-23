from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_table(doc,data):
    # 添加表格
    table = doc.add_table(rows=len(data), cols=len(data[0]))

    # 填充表格数据
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)

    # 设置表格样式
    table.style = 'Table Grid'
    set_table_style(table)

def create_table_by_oxml(doc,oxml):
    tblPr = oxml.find(qn("w:tblPr"))
    # 表格居中方式
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        oxml.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(6*1440))

    # 边框设置
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        existing = tblBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tblBorders.remove(existing)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tblBorders.append(element)

    
    for tc in oxml.findall(".//" + qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        # 背景色清除
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            tcPr.remove(shd)
        # 删除单元格内边距
        # old_tcMar = tcPr.find(qn("w:tcMar"))
        # if old_tcMar is not None:
        #     tcPr.remove(old_tcMar)
        # 垂直居中
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)
        # 清除文字两侧多余空格和缩进
        # for p in tc.findall(".//" + qn("w:p")):

    for run in oxml.findall(".//" + qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run.insert(0, rPr)
        sz = rPr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rPr.append(sz)
        sz.set(qn("w:val"), str(10*2))
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")


    for p in oxml.findall(".//" + qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        # 删除旧的缩进
        old_ind = pPr.find(qn("w:ind"))
        if old_ind is not None:
            pPr.remove(old_ind)

        # 表格水平居中
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

        # 设置表格中文字段落间距
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:before"), "0")

    table = doc.add_table(rows=1, cols=1)
    parent = table._element.getparent()
    parent.replace(table._element, oxml)
    set_table_style(table)

    

def set_table_style(table):
    # 自动调整表格宽度
    table.autofit
    
    tbl = table._element
    tblPr = tbl.tblPr or OxmlElement("w:tblPr")
    tblBorders = tblPr.find(qn("w:tblBorders")) or OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        tblBorders.append(elem)
    tblPr.append(tblBorders)

    for row in table.rows:
        for cell in row.cells:
            # 垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 水平居中
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10)  # 10 磅字号
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


