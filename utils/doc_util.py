from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document


def add_one_level_header(target_doc, header):
    # 一级标题
    heading_paragraph = target_doc.add_paragraph()
    heading_run = heading_paragraph.add_run(header)
    heading_run.bold = True
    heading_run.font.name = "宋体"
    heading_run.font.size = Pt(12)  # 小四对应 12 磅
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 去除行间距
    fmt = heading_paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_content(doc, content):
    # ==================== 正文文字 ====================
    content_paragraph = doc.add_paragraph()
    fmt = content_paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Pt(24)

    if isinstance(content, list):
        for con in content:
            if isinstance(con, dict):
                content_run = content_paragraph.add_run(con.get("content"))
                content_run.bold = con.get("bold", False)
            else:
                content_run = content_paragraph.add_run(str(con))
            # 相同样式
            content_run.font.name = "宋体"
            content_run.font.size = Pt(12)
            content_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        return

    content_run = content_paragraph.add_run(str(content))
    content_run.font.name = "宋体"
    content_run.font.size = Pt(12)
    content_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_cover_page(target_doc):
    """
    生成封面
    """
    target_doc.add_paragraph('')

    # 标题1
    p1 = target_doc.add_paragraph('江苏苏商银行')
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = p1.runs[0]
    run1.bold = True
    run1.font.name = '宋体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run1.font.size = Pt(22)  # 小二

    # 标题2
    p2 = target_doc.add_paragraph('信用债调查报告')
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.runs[0]
    run2.bold = True
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(22)

    # 标题后空行
    for _ in range(5):
        target_doc.add_paragraph('')

    # 字段左缩进
    left_indent_pt = Pt(110)

    labels = [
        '授信客户：',
        '申报部门：',
        '主办调查人：',
        '协办调查人：',
        '联系电话：',
        '部门负责人：'
    ]

    # 为每个标签单独设置下划线空格数（微调以对齐）
    underline_space_counts = [20, 20, 16, 16, 20, 16]

    for i, label in enumerate(labels):
        p = target_doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = left_indent_pt
        p.space_after = Pt(18)

        # 1. 写入标签文字
        run_label = p.add_run(label)
        run_label.font.name = '宋体'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_label.font.size = Pt(18)  # 小三

        # 2. 处理下划线内容逻辑
        if label == '申报部门：':
            content_text = "金融市场部" + (" " * (underline_space_counts[i] - 11)) + "\ufeff"
        else:
            content_text = (" " * underline_space_counts[i]) + "\ufeff"

        run_content = p.add_run(content_text)
        run_content.font.name = '宋体'
        run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_content.font.size = Pt(18)
        run_content.underline = WD_UNDERLINE.SINGLE

    for _ in range(3):
        target_doc.add_paragraph('')

    # 填报时间
    date_p = target_doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    date_p.paragraph_format.left_indent = left_indent_pt

    run_label = date_p.add_run('填报时间：')
    run_label.font.name = '宋体'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_label.font.size = Pt(18)

    # 年
    run_year = date_p.add_run('2026' + '\ufeff')
    run_year.font.name = '宋体'
    run_year._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_year.font.size = Pt(18)
    run_year.underline = WD_UNDERLINE.SINGLE

    run_y = date_p.add_run(' 年 ')
    run_y.font.name = '宋体'
    run_y.font.size = Pt(18)

    # 月
    run_month = date_p.add_run('  ' + '\ufeff')
    run_month.font.name = '宋体'
    run_month._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_month.font.size = Pt(18)
    run_month.underline = WD_UNDERLINE.SINGLE

    run_m = date_p.add_run(' 月 ')
    run_m.font.name = '宋体'
    run_m.font.size = Pt(18)

    # 日
    run_day = date_p.add_run('  ' + '\ufeff')
    run_day.font.name = '宋体'
    run_day._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_day.font.size = Pt(18)
    run_day.underline = WD_UNDERLINE.SINGLE

    run_d = date_p.add_run(' 日')
    run_d.font.name = '宋体'
    run_d.font.size = Pt(18)

    for _ in range(2):
        target_doc.add_paragraph('')

    # 分页符
    target_doc.add_page_break()




if __name__ == "__main__":
    # 生成封面
    doc = Document()
    add_cover_page(doc)
    doc.save('test_credit_report_cover.docx')
    print("测试文档已生成：test_credit_report_cover.docx（仅包含封面）")

