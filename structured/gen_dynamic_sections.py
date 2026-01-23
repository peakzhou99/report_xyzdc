import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# 设置项目路径
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

# 引入项目内部模块
from structured.fetch_structured_report import YJTCompanyDataFetcher
from utils.logger_util import get_logger
from utils.table_util import set_table_style

logger = get_logger()


class ComprehensiveReportGenerator:
    def __init__(self, company_name: str):
        self.company_name = company_name
        self.fetcher = YJTCompanyDataFetcher(company_name=company_name)
        logger.info(f"正在获取【{company_name}】的所有结构化数据...")
        # 一次性获取所有数据
        self.data = self.fetcher.get_report_data()

    # ==========================
    # 通用工具方法
    # ==========================

    def _set_cell_style_simple(self, cell, text, align="center", bold=False, font_size=10):
        """
        通用单元格样式设置（用于动态章节）
        """
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(text) if text is not None else "")

        # 字体设置
        run.font.name = "宋体"
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.bold = bold

        # 对齐方式
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_section_title(self, doc, title):
        """添加章节标题"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(10)

    def _create_compliance_table_structure(self, doc):
        """构造准入条件表的骨架"""
        # 1. 构造硬编码的骨架数据
        raw_rows = [
            ["", "我行标准", "申请人指标说明", "是否符合"],
            ["评级要求\n（满足其一即可）", "1、主体评级不低于AA+，或担保人不低于AA+，或债项评级（如有）不低于AA+；", "", ""],
            ["",
             "2、江苏省内13个地级市的市级平台，或江苏省外一般预算收入超过150亿元的市级平台，主体含担保人及债项评级（如有）可放宽至AA。",
             "", ""],
            ["", "3、区县级平台，且一般预算收入超过20亿元，主体含担保人及债项评级（如有）可放宽至AA；", "", ""],
            ["", "4、省级平台，发行人或担保人或债项评级（如有）不低于AA。", "", ""],
            ["层级要求",
             "江苏省外，投资层级限于层级限于省级、地市级、市属区级、一般预算收入超过20亿元区县级。省外准入区域，采用白名单制，具体区域以名单为准，不定期更新，审慎扩大将我行禁入区域、网红区域、多次出现非标违约的区域纳入白名单。",
             "", ""],
            ["主体资质\n（须全部满足）", "发行人或担保人资产规模不得低于100亿", "", ""],
            ["", "发行人及担保人资产负债率不高于80%", "", ""],
            ["",
             "如发行人及担保人均为市属区级、县级平台，且主体评级均为AA级的，省外发行人或担保人须在本级财政所属区域内净资产规模排名前五，前五包括AA及以上平台；省内不限。",
             "", ""],
            ["禁入区域",
             "江苏省外，原则上天津市、辽宁省、吉林省、黑龙江省、海南省、贵州省、云南省、甘肃省、青海省、内蒙古自治区、广西壮族自治区、西藏自治区、宁夏回族自治区、新疆维吾尔自治区不予准入；重庆市审慎控制业务余额。",
             "", ""],
            ["审慎介入区域",
             "区县级工所属地市级宽口径债务率超800%的、地市范围内近三年出现过多次实质性债务违约的等，除有较强增信外（如省级AAA融担公司担保），审慎介入。",
             "", ""],
            ["其他禁入条件", "主体及担保人评级和债项评级（如有）近一年未被下调评级，且展望不得为负面", "", ""],
            ["",
             "近三年地方所在政府或下属平台公司出现过债务违约（实质性城投非标债务逾期）的区域（区县级范围），存量债务清理甄别中存在恶意改变债权债务关系的区域，或是政府在PPP、政府购买服务或是政府须承担付款责任的其他类型项目中存在重大违约记录导致银行债权无法实现、形成不良的区域；",
             "", ""],
            ["", "我行发布过风险预警、禁入的主体，或出现在交易对手警示名单之列；", "", ""],
            ["管控要求", "区域限额是否超标（限额标准见风险指引，标明该区域目前已批已投金额）", "", ""],
            ["",
             "单户投资额度（发行人或保证人孰高，AAA城投债单户金额不超过8000万元，AA+城投债单户金额不超过6000万元，AA城投债单户金额不超过4000万元，债券单户投资额度不超过单支债券（以债券编号为准）实际发行金额的20%）是否超标；",
             "", ""],
            ["", "我行持有期限不超过2年", "", ""]
        ]

        # 2. 创建表格
        table = doc.add_table(rows=len(raw_rows), cols=len(raw_rows[0]))
        table.style = 'Table Grid'
        table.allow_autofit = False

        # 3. 填充静态文本
        for i, row in enumerate(raw_rows):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(str(value))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run.font.name = "宋体"
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 4. 执行合并 (Indices must match raw_rows structure)
        table.cell(1, 0).merge(table.cell(4, 0))  # 评级要求
        table.cell(6, 0).merge(table.cell(8, 0))  # 主体资质
        table.cell(11, 0).merge(table.cell(13, 0))  # 其他禁入条件
        table.cell(14, 0).merge(table.cell(16, 0))  # 管控要求

        # 5. 应用外部样式工具并调整列宽
        set_table_style(table)

        col_widths = [Pt(70), Pt(275), Pt(95), Pt(60)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                cell = row.cells[idx]
                tcPr = cell._element.get_or_add_tcPr()
                tcW = tcPr.get_or_add_tcW()
                tcW.type = 'dxa'
                tcW.w = int(width.pt * 20)

                # 清理合并产生的空行
                if len(cell.paragraphs) > 1:
                    for extra_p in cell.paragraphs[1:]:
                        if not extra_p.text.strip():
                            p_element = extra_p._element
                            p_element.getparent().remove(p_element)

                # 对齐微调：第二列靠左，其余居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER

        return table

    def _fill_compliance_data(self, table):
        """将获取的结构化数据填充到准入表中"""
        # 映射逻辑：(行索引, 数据源字段)
        mapping = {
            1: self.data.get("评级要求", {}).get("评级要求1", {}),
            2: self.data.get("评级要求", {}).get("评级要求2", {}),
            3: self.data.get("评级要求", {}).get("评级要求3", {}),
            4: self.data.get("评级要求", {}).get("评级要求4", {}),
            5: self.data.get("层级要求", {}),
            6: self.data.get("主体资质1", {}),
            7: self.data.get("主体资质2", {}),
            8: self.data.get("主体资质3", {}),
            9: self.data.get("禁入区域", {}),
            11: self.data.get("其他禁入条件", {}),
        }

        for row_idx, data_item in mapping.items():
            if not data_item:
                continue

            desc = data_item.get("申请人指标说明", "")
            result = data_item.get("符合/不符合", "")

            # 填充第2列（说明）和第3列（结果）
            cell_desc = table.cell(row_idx, 2)
            cell_desc.text = str(desc)
            cell_result = table.cell(row_idx, 3)
            cell_result.text = str(result)

            # 重新应用字体样式（text赋值会重置样式）
            for cell in [cell_desc, cell_result]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_compliance_section(self, doc):
        """生成完整的准入章节"""
        self._add_section_title(doc, "准入条件审查表")
        table = self._create_compliance_table_structure(doc)
        self._fill_compliance_data(table)
        doc.add_page_break()  # 准入表后换页


    def _create_basic_info_table(self, doc):
        """一、公司基本信息"""
        self._add_section_title(doc, "一、公司基本信息")
        info = self.data.get("公司基本信息", {})

        fields = [
            ("注册名称", info.get("COMPNAME")),
            ("法定代表人", info.get("LEGREP")),
            ("注册资本", f"{info.get('REGCAPITAL', '')} 万元"),
            ("设立日期", info.get("FOUNDDATE")),
            ("统一社会信用代码", info.get("unified_credit_code")),
            ("公司住所", info.get("REGADDR")),
            ("办公地址", info.get("OFFICEADDR")),
            ("邮政编码", info.get("OFFICEZIPCODE")),
            ("电话", info.get("COMPTEL")),
            ("传真", info.get("COMPFAX")),
            ("经营范围", info.get("BIZSCOPE"))
        ]

        table = doc.add_table(rows=len(fields), cols=2)
        table.style = 'Table Grid'

        for i, (label, value) in enumerate(fields):
            self._set_cell_style_simple(table.cell(i, 0), label, align="left", bold=True)
            self._set_cell_style_simple(table.cell(i, 1), value, align="left")

        # 简单设置列宽
        for row in table.rows:
            row.cells[0].width = Pt(100)
            row.cells[1].width = Pt(350)

    def _create_shareholder_table(self, doc):
        """二、股东信息"""
        self._add_section_title(doc, "二、股东信息")
        shareholders = self.data.get("股东信息", [])

        if not shareholders:
            doc.add_paragraph("暂无股东信息")
            return

        headers = ["股东名称", "股东内码", "持股数量", "持股比例"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            self._set_cell_style_simple(table.rows[0].cells[i], header, bold=True)

        for item in shareholders:
            row_cells = table.add_row().cells
            self._set_cell_style_simple(row_cells[0], item.get("SHHOLDERNAME"))
            self._set_cell_style_simple(row_cells[1], item.get("SHHOLDERSECODE"))
            self._set_cell_style_simple(row_cells[2], item.get("HOLDERAMT"))
            self._set_cell_style_simple(row_cells[3], f"{item.get('HOLDERRTO', '')}%" if item.get('HOLDERRTO') else "")

    def _create_rating_table(self, doc):
        """三、主体评级列表"""
        self._add_section_title(doc, "三、主体评级列表")
        ratings = self.data.get("主体评级列表", [])

        if not ratings:
            doc.add_paragraph("暂无评级信息")
            return

        headers = ["评级公司", "级别日期", "级别", "展望", "披露日期", "有效截止日", "级别对象"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            self._set_cell_style_simple(table.rows[0].cells[i], header, bold=True)

        for item in ratings:
            row_cells = table.add_row().cells
            self._set_cell_style_simple(row_cells[0], item.get("RATECOMNAME"))
            self._set_cell_style_simple(row_cells[1], item.get("PUBLISHDATE"))
            self._set_cell_style_simple(row_cells[2], item.get("CREDITRATE"))
            self._set_cell_style_simple(row_cells[3], item.get("EXPTRATING_value"))
            self._set_cell_style_simple(row_cells[4], item.get("DECLAREDATE"))
            self._set_cell_style_simple(row_cells[5], item.get("CREDITRATEENDDATE"))
            self._set_cell_style_simple(row_cells[6], item.get("COMTYPE_value"))

    def _create_platform_score_table(self, doc):
        """四、区域排名信息"""
        self._add_section_title(doc, "四、区域排名信息")
        scores = self.data.get("区域排名信息", [])

        if not scores:
            doc.add_paragraph("暂无区域排名信息")
            return

        headers = ["公司名称", "区域", "城投评分", "省内排名", "主体评级", "债券余额"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            self._set_cell_style_simple(table.rows[0].cells[i], header, bold=True)

        for item in scores:
            row_cells = table.add_row().cells
            self._set_cell_style_simple(row_cells[0], item.get("itname"))
            self._set_cell_style_simple(row_cells[1], item.get("regionname"))
            self._set_cell_style_simple(row_cells[2], item.get("score_all"))
            self._set_cell_style_simple(row_cells[3], item.get("rank"))
            self._set_cell_style_simple(row_cells[4], item.get("credit_rate"))
            self._set_cell_style_simple(row_cells[5], item.get("bondbalance"))

    def _create_spreads_table(self, doc):
        """六、利差信息"""
        self._add_section_title(doc, "六、利差信息")
        spreads = self.data.get("利差信息", [])

        desc_p = doc.add_paragraph()
        run = desc_p.add_run(f"{self.company_name} 近七天平均信用利差如下：")
        run.font.name = "宋体"
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        if not spreads:
            doc.add_paragraph("暂无利差信息")
            return

        headers = ["交易日期", "当前利差(BP)"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            self._set_cell_style_simple(table.rows[0].cells[i], header, bold=True)

        for item in spreads:
            row_cells = table.add_row().cells
            self._set_cell_style_simple(row_cells[0], item.get("tradedate"))
            self._set_cell_style_simple(row_cells[1], item.get("spread"))

    def _create_registrations_table(self, doc):
        """七、注册批复全景"""
        self._add_section_title(doc, "七、注册批复全景")
        regs = self.data.get("注册批复全景", [])

        if not regs:
            doc.add_paragraph("暂无注册批复信息")
            return

        headers = ["项目名称", "批复场所", "品种", "注册中额度", "最新注册时间", "最新注册状态"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            self._set_cell_style_simple(table.rows[0].cells[i], header, bold=True)

        for item in regs:
            row_cells = table.add_row().cells
            self._set_cell_style_simple(row_cells[0], item.get("project_name"))
            self._set_cell_style_simple(row_cells[1], item.get("approval_location"))
            self._set_cell_style_simple(row_cells[2], item.get("project_type_name"))
            self._set_cell_style_simple(row_cells[3], item.get("register_amount"))
            self._set_cell_style_simple(row_cells[4], item.get("process_date"))
            self._set_cell_style_simple(row_cells[5], item.get("process_type_name"))

    def generate(self, output_path: str):
        """生成整合报告"""
        logger.info(f"开始生成【{self.company_name}】的综合报告...")
        doc = Document()

        # 1. 准入条件表（作为首页或前言）
        self._add_compliance_section(doc)

        # 2. 动态详细章节
        self._create_basic_info_table(doc)
        doc.add_paragraph()

        self._create_shareholder_table(doc)
        doc.add_paragraph()

        self._create_rating_table(doc)
        doc.add_paragraph()

        self._create_platform_score_table(doc)
        doc.add_paragraph()

        self._create_spreads_table(doc)
        doc.add_paragraph()

        self._create_registrations_table(doc)

        doc.save(output_path)
        logger.info(f"综合报告已生成至: {output_path}")


if __name__ == "__main__":
    target_company = "泰安泰山城乡建设发展有限公司"
    save_file = f"{target_company}_综合信用报告.docx"

    generator = ComprehensiveReportGenerator(target_company)
    generator.generate(save_file)