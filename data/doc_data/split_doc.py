import difflib
import asyncio
import os, re
import time
import pandas as pd
import json
import zipfile
import uuid
from typing import Dict

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from utils.string_util import extract_date, extract_unit
from utils.log_utils import get_logger
from analyze.private_apply import PrivateApply
from data.doc_data.doc_table_processor import TableProcessor

logger = get_logger()
semaphore = asyncio.Semaphore(10)

class SplitDoc():
    """
    文档内容是大模型进行解析的依据，由于篇幅较长，为快速定位及解析内容的准确度，故将文章按照章节目录拆分；
    在拆分的过程中发现，目录标题存在自动编码情况，对此无法识别其编号；
    通过多次尝试，发现转为pdf文档后，能够将自动编号解析出来，所以首先转为pdf文档，之后解析pdf内容，根据段落长度以及是否含有对应标题前缀，来确认标题范围；
    不过，通过上面确认的标题范围，存在很多脏数据，需要再次通过word中的标题做进一步的清洗；
    """

    def __init__(self, doc_path, company_name):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.private_apply_instance = PrivateApply()
        self.table_data_set = [company_name]
        # 类似标题前缀
        self.hz_seq_2dlist = [[]]
        self.hz_seq_list = []
        self.level_seqs = []
        oth_level_seqs_list = []
        self._init_hz_seq()
        self.paragraphs = self.parse_doc()
        self.table_datas = {}
        self.table_oxmls = {}
        self.article_datas = {}
        self.company_info = {}
        self.fuzhai_total = []

        self.menus = [title.get("content") for title in self.parse_doc_title()]
        self.table_keywords_exclude = ["发行人受限货币资金", "营业成本具体构成情况如下", "经营性(其他应收|分类统计)", "政府补助的应收款项"]  # 有息负债分类情况

        self.table_keywords = [
            "合并资产", "资产负债",
            "合并利润", "利润表", "损益",
            "合并现金", "现金流量",
            "外担保", "对外提供担保", "担保事项",
            "应收", "受限", "资产限制情况", "受到限制", "抵押", "质押",
            "有息", "信托", "委托贷款",
            "授信", "营业", "收入"
        ]

        self.table_type_list = ["资产负债表", "利润表", "应收账款", "其他应收款", "受限资产", "有息负债", "授信情况", "营业收入", "现金流", "对外担保", "保证人基本情况"]

        self.target_table_descs = {
            "资产负债表": "资产负债表:必须是时间最新的表(根据表格标题和时间序列判断)，必须是发行人合并资产负债表，必须含有资产类项目和负债类项目并且有他们的合计，时间序列至少有2个不同",
            "利润表": "利润表:必须是时间最新的表(根据表格标题和时间序列判断)，必须是发行人合并利润表，表格标题可能含有'利润'、'损益'等字样，必须含有营业总收入、营业利润、利润总额、净利润等项目，时间序列至少有2个不同；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "应收账款": "应收账款:必须是时间最新的表(根据表格标题和时间序列判断)，表格标题可能含有“应收账款”，若同时含有“主要”或“前五”等，需要优先识别；表格头含有单位名称(如项目名称)、金额和占比等字段或相似字段，可能含有账龄和款项性质；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "其他应收款": "其他应收款:必须是时间最新的表(根据表格标题和时间序列判断)，表格头含有债务人（如单位名称）、期末余额（如金额）和占比等字段，可能含有账龄和款项性质；优先取包含前五名应收款信息的表；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "受限资产": "受限资产:必须是时间最新的表(根据表格标题和时间序列判断)，表格头含项目、受限金额（如账面价值）和受限制原因等，可能含有占比；表格内容多数含有货币资金和存货等；如果有相似表优先取表头中包含受限原因的表；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "有息负债": "有息负债:必须是时间最新的表(根据表格标题和时间序列判断)，标题可能含有期限结构等表述；重点关注一年内的债务欠款，可能表述为'1年以内（含1年）'、'1年内到期'等；若包含“银行贷款”、“债券融资”和“非标融资”等项目，时间序列至少有2个不同；如果有相似表，优先取最新时间的表，其次取表头或关键词中明确包含1年内到期借款的表；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "授信情况": "授信情况:必须是时间最新的表(根据表格标题和时间序列判断)，表格头含银行、授信额度、已使用额度、未使用额度等；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "营业收入": "营业收入:必须是时间最新的表(根据表格标题和时间序列判断)，表头或关键词中含有项目（如业务板块）以及季度或年度的金额和占比；时间序列至少有2个不同；优先取包含主营业务各板块收入信息的表，不选表头或关键词包含营业成本的表；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "现金流": "现金流:必须是时间最新的表(根据表格标题和时间序列判断)，必须是发行人合并现金流，重点关注短期借款；时间序列至少有2个不同；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "对外担保": "对外担保：必须时间最新的表，表头含有被担保人(方)、担保余额（金额）和担保期限等；若表格标题含有主体信息，则主体必须为发行人，若主体为担保人，则必须排除该表格",
            "保证人基本情况": "保证人情况:表格标题可能含有保证人（如担保人），关键词含有注册地址、法定代表人和注册资本等",
        }

    def _init_hz_seq(self):
        hz0_seq = ["募集说明书", "声明", "重大事项", "重要提示", "释义", "目录"]
        hz1_seq = ["第一", "第二", "第三", "第四", "第五", "第六", "第七", "第八", "第九", "第十", "第十一", "第十二",
                   "第十三", "第十四", "第十五", "第十六", "第十七", "第十八", "第十九", "第二十", "第二十一",
                   "第二十二", "第二十三", "第二十四", "第二十五", "第二十六", "第二十七", "第二十八", "第二十九"]
        hz2_seq = ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "十一、", "十二、", "十三、",
                   "十四、", "十五、", "十六、", "十七、", "十八、", "十九、", "二十、", "二十一、", "二十二、", "二十三、",
                   "二十四、", "二十五、", "二十六、", "二十七、", "二十八、", "二十九"]
        hz3_seq = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）", "（十一）", "（十二）",
                   "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）", "（二十一）",
                   "（二十二）", "（二十三）", "（二十四）", "（二十五）", "（二十六）", "（二十七）", "（二十八）", "（二十九）"]
        hz4_seq = ["1、", "2、", "3、", "4、", "5、", "6、", "7、", "8、", "9、", "10、", "11、", "12、", "13、", "14、", "15、",
                   "16、", "17、", "18、", "19、", "20、", "21、", "22、", "23、", "24、", "25、", "26、", "27、", "28、", "29、"]
        hz5_seq = ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.", "13.", "14.", "15.",
                   "16.", "17.", "18.", "19.", "20.", "21.", "22.", "23.", "24.", "25.", "26.", "27.", "28.", "29."]
        hz6_seq = ["（1）", "（2）", "（3）", "（4）", "（5）", "（6）", "（7）", "（8）", "（9）", "（10）", "（11）", "（12）", "（13）",
                   "（14）", "（15）", "（16）", "（17）", "（18）", "（19）", "（20）", "（21）", "（22）", "（23）", "（24）", "（25）",
                   "（26）", "（27）", "（28）", "（29）"]
        hz7_seq = ["1）", "2）", "3）", "4）", "5）", "6）", "7）", "8）", "9）", "10）", "11）", "12）", "13）", "14）", "15）",
                   "16）", "17）", "18）", "19）", "20）", "21）", "22）", "23）", "24）", "25）", "26）", "27）", "28）", "29）"]
        hz8_seq = ["(一)", "(二)", "(三)", "(四)", "(五)", "(六)", "(七)", "(八)", "(九)", "(十)", "(十一)", "(十二)",
                   "(十三)", "(十四)", "(十五)", "(十六)", "(十七)", "(十八)", "(十九)", "(二十)", "(二十一)",
                   "(二十二)", "(二十三)", "(二十四)", "(二十五)", "(二十六)", "(二十七)", "(二十八)", "(二十九)"]

        self.hz_seq_2dlist = [hz0_seq, hz1_seq, hz2_seq, hz3_seq, hz4_seq, hz5_seq, hz6_seq, hz7_seq]
        self.hz_seq_list = hz0_seq + hz1_seq + hz2_seq + hz3_seq + hz4_seq + hz5_seq + hz6_seq + hz7_seq + hz8_seq
        self.level_seqs = [hz0_seq + hz1_seq]
        self.first_level_hz_seq = hz0_seq + hz1_seq
        self.oth_level_seqs_list = [hz2_seq, hz3_seq, hz4_seq, hz5_seq, hz6_seq, hz7_seq]
        self.other_level_hz_seq = hz2_seq + hz3_seq + hz4_seq + hz5_seq + hz6_seq + hz7_seq

    def _is_first_level_title(self, text):
        """判断是否为一级标题"""
        if text is None:
            return False
        if len(text) > 50:
            return False
        # 判断是否为目录内容
        if text.startswith(tuple(self.hz_seq_list)) and bool(re.search(r'[0-9]+$', text)):
            return False
        if text.startswith(tuple(self.first_level_hz_seq)):
            return True
        if text.startswith(tuple(self.other_level_hz_seq)):
            return False
        if any(text.startswith(_hz_seq) for _hz_seq in self.first_level_hz_seq):
            return True

    def _is_title(self, paragraph, context):
        _is_potential_title = self._is_potential_title(paragraph)
        if _is_potential_title is None:
            return self.is_title_by_llm(paragraph.get("content"), context)
        else:
            return _is_potential_title

    def _is_potential_title(self,paragraph):
        """强规则"""
        text = paragraph.get("content")
        style = paragraph.get("style")
        align = style.get("align") if style is not None else None
        if len(text) > 50:
            return False
        # 特定字符开头不是标题
        if bool(re.search(r"^[①②③④⑤⑥a-zA-Z](\s)?[\\.、）)]?.*", text)):
            return False
        # 特定字符结尾不是标题
        if bool(re.search(r'[\?\!\.\;\、\？\！\。\；\/]+$', text)):
            return False
        # 含有冒号且含有特殊字符不是标题
        if bool(re.search(r"(名称|法定代表人|所属行业|实缴资本|设立日期|统一社会信用代码|数据来源|住所|办公地址|传真|邮政编码|联系|电话|负责人|联系人|包括|上述|如下|以下|下列|介绍|其中|借|贷|注|单位)+.*?[\\:：]+", text)):
            return False
        # 包含特殊字符都不是标题
        if bool(re.search(r'[【】—]', text)):
            return False
        # 居右的文本不是标题
        if align is not None and align.lower() =='right':
            return False
        # 除了一级标题居中，其他标题都是居左
        if align is not None and align.lower() =='center' and (not self._is_first_level_title(text)):
            return False
        # 排除页眉页脚
        if bool(re.search(r'^[0-9]+$', text)) or bool(re.search(r'.{6,}募集说明书$', text)):
            return False
        # 剔除目录内容：前部分是标题，结尾是页码
        if text.startswith(tuple(self.hz_seq_list)) and bool(re.search(r'[0-9]+$', text)):
            return False
        elif text.startswith(tuple(self.hz_seq_list)) and not bool(re.search(r'[\:\,\?\!\.\;\：\，\、\？\！\。\；\/]+$', text)):
            return True

    def is_title(self, text: str, context: str = None):
        """
        1. 规则判断是否为标题
        2. 大模型判断是否为标题
        """
        # 标题内容不能太长
        if len(text) > 50:
            return False

        # 除冒号外的标点符号结尾都不认为是标题
        if bool(re.search(r'[\?\!\.\;\、\？\！\。\；\/]+$', text)):
            return False
        # 包含特殊字符都不认为是标题
        if bool(re.search(r'(【|】|签字(）)*：|,|，)+', text)) or bool(re.search(r'：', text[:-1])):
            return False
        # 以特殊字符开头也不认为标题
        if text.startswith(tuple(["单位：", "表", "图", "目录", "注：", "年月日", "_"])):
            return False
        # 排除页眉页脚
        if bool(re.search(r'^[0-9]+$', text)) or bool(re.search(r'.{6,}募集说明书$', text)):
            return False
        # 判断是否为目录内容
        if text.startswith(tuple(self.hz_seq_list)) and bool(re.search(r'[0-9]+$', text)):
            return False
        elif text.startswith(tuple(self.hz_seq_list)) and not self.ends_with_punctuation(text):
            return True
        # 利用大模型判断是否为标题
        if context and self.is_title_by_llm(text, context):
            return True

    def extract_paragraphs(self):
        """
        按照标题拆分文章内容，存储以K-V方式内存存储，Key:标题面包屑导航, Value:该标题下内容
        """
        start_time = time.time()
        logger.debug("开始抽取文档目录")
        # 获取标题列表
        doc_title_list = self.parse_doc_title()
        # 设置标题层级
        menu_list = self.get_menu_list(doc_title_list)
        # 面包屑导航
        menu_breakcrumb = self.get_menu_breakcrumb(menu_list)
        # 拆分文章
        chapter_dict = self.split_by_menu_with_tbl(menu_breakcrumb)
        self.article_datas.update(chapter_dict)
        # 发行人和担保人名称提取
        self.company_info = self.company_extractor(chapter_dict)
        logger.debug(f"文档目录抽取结束，耗时： {time.time() - start_time}")

    def extract_tables(self):
        logger.debug("开始抽取表格")
        start_time = time.time()
        # 抽取表格上下文
        table_contexts = self.extract_table_context_with_menu()
        # 关键词过滤表格
        filtered_table_contexts = self.filter_table_by_keywords(table_contexts)
        # 从文档中抽取目标表格
        target_tables = self.extract_target_tables_from_doc(filtered_table_contexts)
        # 格式处理
        self.table_data_format(target_tables)
        # 输出到集合中
        self.create_table_data_set(target_tables)
        # print(self.table_datas.keys())
        logger.debug(f"表格解析结束，总耗时： {time.time() - start_time}")

    def parse_doc(self):
        start_time = time.time()
        doc = Document(self.doc_path)
        doc_elem = parse_xml(doc.part.element.xml)
        namespaces = doc_elem.nsmap
        self.namespaces = namespaces
        body_elem = doc_elem.find("w:body", namespaces)
        # 文档段落
        paragraphs = []
        for para_elem in body_elem.getchildren():
            if para_elem.tag == qn("w:p"):
                # 过滤文本为空情况
                if (para_elem.text is None or para_elem.text.strip().replace(" ", "") == '' or para_elem.text.strip().replace("	", "") == ''):
                    continue
                text = para_elem.text.strip().replace(" ", "").replace("	", "")
                # 文档样式
                style_el = para_elem.find(".//w:pPr/w:pStyle", namespaces)
                style_id = style_el.get(qn('w:val')) if style_el is not None else None
                outline_el = para_elem.find(".//w:pPr/w:outlineLvl", namespaces)
                level = int(outline_el.get(qn("w:val"))) if outline_el is not None else None
                numId_el = para_elem.find(".//w:pPr/w:numPr/w:numId", namespaces)
                numId = numId_el.get(qn('w:val')) if numId_el is not None else None
                ilvl_el = para_elem.find(".//w:pPr/w:numPr/w:ilvl", namespaces)
                ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else None
                align_el = para_elem.find(".//w:pPr/w:jc", namespaces)
                align = align_el.get(qn('w:val')) if align_el is not None else None
                style = {"styleId": style_id, "outlineLvl": level, "numId": numId, "ilvl": ilvl, "align": align}
                paragraphs.append({"type": "text", "content": text, "oxml": para_elem, "level": None, "style": style})
            elif para_elem.tag == qn("w:tbl"):
                tab_datas = self.parse_oxml_to_2list(para_elem, namespaces)
                style = {"styleId": None, "outlineLvl": None, "numId": None, "ilvl": None, "align": None}
                paragraphs.append({"type": "table", "content": tab_datas, "oxml": para_elem, "level": None, "style": style})
        logger.debug(f"文档解析结束，总耗时： {time.time() - start_time}")
        return paragraphs

    def is_title_by_llm(self, text: str, context: str = None):
        title_messages = [{
            "role": "user",
            "content": f"""/no_think
            你是金融文档分析专家，根据**段落上下文**，判断所提供的**文本数据**是否为标题，不做任何解释，只以JSON格式返回是或否。
            【段落上下文】
            {context}
            【文本数据】
            {text}
            【注意事项】
            1. **段落上下文**是列表格式，列表元素代表一个标题、段落、表名称或目录内容；
            2. **文本数据**是**段落上下文**的一个元素；
            3. 目录内容虽然符合标题的要求，但结尾会带页码(阿拉伯数字或罗马数字)，结果输出否；
            4. **文本数据**前后连续出现带标题序号的单句，虽然符合标题要求,但单句之间没有内容文本，不符合标题-内容结构，结果输出否；
            5. 表格标题后文中紧跟着表格单位，不属于标题，结果输出否；
            6. 表名称中一般带“表：”，结果输出否；
            7. **文本数据**为标题时，可能没有文字序号或数字序号；
            【结果输出】
            {{"结果": ""}}
            【结果输出】
            {{"结果": "否"}}
            """
            }]
        result = self.private_apply_instance.generate(title_messages)
        if "是" in result:
            return True
        else:
            return False

    def company_extractor(self, sections):
        # 抽取发行人和担保人公司名称
        basic_infos = sections.get("募集信息")[:30000]
        messages = [{
            "role": "user",
            "content": f"""/no_think
                            你是金融文档分析专家，请从“募集要素信息”中提取关键信息，并以JSON格式输出：
                            1. 提取**发行人公司名称**；
                            2. 判断是否存在增信情况（如担保情况）；
                            3. 若存在增信情况，提取所有**担保人公司名称**（可能多个）
                            4. 结果以JSON格式输出；
                            【募集要素信息】
                            {basic_infos}
                            【结果输出】
                            {{"发行人": "", "是否存在增信": "是" 或 "否", "担保人": []}}
                            """
            }]
        result = self.private_apply_instance.generate(messages)
        return json.loads(result)

    def ends_with_punctuation(self, text: str):
        """判断结尾标点符号"""
        return bool(re.search(r'[\:\,\?\!\.\;\：\，\、\？\！\。\；\/]+$', text))

    def get_title_level(self, title):
        # 若存在之前层级中，则直接返回层级
        for index, item in enumerate(reversed(self.level_seqs)):
            if title.startswith(tuple(item)) or title.endswith(tuple(item)):
                level = len(self.level_seqs) - index
                self.level_seqs = self.level_seqs[:level]
                return level
        # 若不在之前层级中，则追加新层级
        for item in self.oth_level_seqs_list:
            if title.startswith(tuple(item)) and title.startswith(item[0]):
                self.level_seqs.append(item)
                return len(self.level_seqs)
        return -1

    def get_menu_list(self, titles):
        # 提取大纲
        outlines = [
            {"index": idx, "content": para.get("content"), "level": para['level']}
            for idx, para in enumerate(self.paragraphs)
            if para['level'] is not None and para['level'] < 10 and len(para['content']) < 80
        ]
        outline_indices = [outline["index"] for outline in outlines]
        # 修复层级
        outlines = self._update_level(outlines, titles)
        # 合并标题
        all_titles = outlines + [
            {"index": title["index"], "content": title.get("content"), "level": 9}
            for title in titles
            if title["index"] not in outline_indices
        ]
        # 标题排序
        all_titles.sort(key=lambda x: x["index"])
        _end_idx = next((i for i, item in enumerate(all_titles) if "声明" in item["content"]), -1)
        all_titles = self._remove_between(all_titles, end=_end_idx - 1)
        # 补充募集概要信息
        all_titles.insert(0, {"index": 0, "content": "募集信息", "level": 0})
        return all_titles

    def _update_level(self, outlines, titles):
        # 构建一级标题字典
        first_level_map = {title["index"]: title for title in titles if title["is_first_level"]}
        # 提取大纲
        tmp_outlines = []
        for outline in outlines:
            _index = outline.get("index")
            if _index in first_level_map and first_level_map.get(_index).get("is_first_level"):
                outline["level"] = 0
            tmp_outlines.append(outline)
        # 大纲级别规范化
        outlines_size = len(tmp_outlines)
        curr_title_idx = -1
        for tmp_idx, tmp_outline in enumerate(tmp_outlines):
            pre_title_idx = curr_title_idx
            # 判断一级标题
            if tmp_outline.get("level") == 0 or tmp_outline.get("index") in first_level_map:
                curr_title_idx = tmp_idx
            # 判断最后一个元素
            elif tmp_idx + 1 == outlines_size:
                curr_title_idx = outlines_size
            # 遍历元素
            if (pre_title_idx >= 0 and pre_title_idx != curr_title_idx
                    and any(inn_outline.get("level") == 0 for inn_outline in
                            tmp_outlines[pre_title_idx + 1:curr_title_idx])):
                for upd_idx in range(pre_title_idx + 1, curr_title_idx):
                    tmp_outlines[upd_idx]["level"] = tmp_outlines[upd_idx]["level"] + 1
        return tmp_outlines

    def _remove_between(self, items, start=None, end=None):
        if (start is None or start < 0) and (end is None or end < 0):
            return items
        _start_idx = 0
        _end_idx = len(items)
        if start is not None and start >= 0:
            _start_idx = start
        if end is not None and end >= 0 and end < len(items):
            _end_idx = end
        del items[_start_idx:_end_idx + 1]
        return items

    def get_menu_breakcrumb(self, menu_list):
        """面包屑导航"""
        title_link = []
        breakcrumb_titles = []
        for _title in menu_list:
            curr_level = _title['level']
            if not title_link:
                title_link = [_title]
            elif curr_level == 0:
                title_link = [_title]
            elif curr_level > title_link[-1]["level"]:
                title_link.append(_title)
            else:
                title_link = self._pop_until_lower_level(title_link, curr_level)
                title_link.append(_title)
            breakcrumb_titles.append({"index": _title["index"], "title": _title['content'], "breakcrumb": '-'.join([inn_menu["content"] for inn_menu in title_link])})
        return breakcrumb_titles

    def _pop_until_lower_level(self, chain: list, lvl: str) -> list:
        while chain and chain[-1]["level"] >= lvl:
            chain.pop()
        return chain

    def parse_oxml_to_2list(self, oxml_table, namespaces):
        """将表格数据转为二维数组数据"""
        curr_table = []
        for row in oxml_table.iterfind("w:tr", namespaces):
            cell_idx = -1
            row_data = []
            for cell in row.iterfind("w:tc", namespaces):
                cell_idx += 1
                vmerge = True if cell.find('.//w:vMerge', namespaces) is not None else False
                text = "".join([para.text for para in cell.iterfind("w:p", namespaces)])
                if vmerge and text == "":
                    if len(curr_table) > 0:
                        text = curr_table[-1][cell_idx]
                    else:
                        text = ""
                # 获取列合并
                grid_span = cell.find('.//w:gridSpan', namespaces)
                if grid_span is not None:
                    grid_span_val = grid_span.val
                    row_data += [text for _ in range(grid_span_val)]
                else:
                    row_data.append(text)
            curr_table.append(row_data)
        return curr_table

    def parse_oxml_to_list_no_header(self, oxml_table, namespaces):
        tmp_result = []
        curr_table = self.parse_oxml_to_2list(oxml_table, namespaces)
        # 不存在，直接返回
        if not curr_table:
            return curr_table
        # 存在，将一行转为一条文本
        for curr_row in curr_table:
            tmp_result.append(":".join(curr_row))
        return tmp_result

    def split_by_menu_with_tbl(self, menu_breakcrumb):
        # 构建标题索引
        breakcrumb_title_map = {breakcrumb_title["index"]: breakcrumb_title for breakcrumb_title in menu_breakcrumb}
        # 标题切分
        section_dict = {}
        curr_breakcrumb_title = {}
        for idx, para in enumerate(self.paragraphs):
            if idx in breakcrumb_title_map:
                curr_breakcrumb_title = breakcrumb_title_map[idx]
            if not curr_breakcrumb_title:
                continue
            if curr_breakcrumb_title["breakcrumb"] in section_dict:
                section_dict.get(curr_breakcrumb_title["breakcrumb"]).append({"index": idx, "type": para.get("type"), "content": para.get("content"), "oxml": para.get("oxml")})
            else:
                section_dict[curr_breakcrumb_title["breakcrumb"]] = [{"index": idx, "type": para.get("type"), "content": para.get("content"), "oxml": para.get("oxml")}]
        return section_dict

    def parse_doc_title(self):
        # 大纲级别
        style_outline = self.build_style_outline_level()
        # 识别自动编码标题
        numbering_map = self.get_numbering_num()
        # 大纲补全
        self.fixed_paragraphs(style_outline, numbering_map)
        # 提取标题
        title_list = self.detect_title()
        return title_list

    def build_style_outline_level(self):
        style_map = {}
        with zipfile.ZipFile(self.doc_path) as docx_zip:
            styles_xml = docx_zip.read("word/styles.xml")
            styles_root = parse_xml(styles_xml)
            styles_ns = styles_root.nsmap
            # 遍历所有样式
            for style in styles_root.findall("w:style", styles_ns):
                style_id = style.get(qn("w:styleId"))
                based_on = style.find("w:basedOn", styles_ns)
                base_id = based_on.get(qn("w:val")) if based_on is not None else None
                if style_id is not None:
                    # 大纲级别
                    outline_lvl_el = style.find(".//w:outlineLvl", styles_ns)
                    level = int(outline_lvl_el.get(qn("w:val"))) if outline_lvl_el is not None else 100
                    # 自动编码
                    numId_el = style.find(".//w:numPr/w:numId", styles_ns)
                    numId = numId_el.get(qn("w:val")) if numId_el is not None else None
                    ilvl_el = style.find(".//w:numPr/w:ilvl", styles_ns)
                    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else None
                    style_map[style_id] = {"outlineLvl": level, "numId": numId, "ilvl": ilvl, "basedOn": base_id}
            # 递归补齐缺失的 level
            for sid in style_map:
                style_map[sid]["outlineLvl"] = self._resolve_level(style_map, sid)
        return style_map

    def _resolve_level(self, style_map: Dict[str, Dict], style_id: str, depth=0) -> int:
        """
        递归获取样式的层级
        """
        if depth > 20:
            return 100
        # #1. 获取当前样式节点样式信息
        style_info = style_map.get(style_id)
        if not style_info:
            return 100
        # #2. 判断是否存在 basedOn
        parent = style_info.get("basedOn")
        if parent is not None:
            _current_level = style_info.get("outlineLvl")
            _parent_level = self._resolve_level(style_map, parent, depth + 1)
            return min(_current_level, _parent_level)
        elif style_info.get("outlineLvl") is not None:
            return style_info["outlineLvl"]
        else:
            return 100

    def fixed_paragraphs(self, style_outline, numbering_map):
        for idx, para_info in enumerate(self.paragraphs):
            # 过滤非文本情况
            if para_info.get("type") == "table":
                continue
            # 直接获取大纲级别
            _level = 100
            style = para_info.get("style")
            if style is not None and style.get("outlineLvl") is not None:
                _level = min(style.get("outlineLvl"), _level)
            elif style is not None and style.get("styleId") is not None:
                style_info = style_outline.get(style.get("styleId"))
                if style_info is not None:
                    _level = min(_level, style_info.get("outlineLvl", 100))
            if style is not None and numbering_map.get((style.get("numId"), style.get("ilvl"))) is not None:
                _level = min(9, _level)
            self.paragraphs[idx]["level"] = _level

    def get_numbering_num(self):
        """
        获取带序号格式的标题
        """
        num_mapping = {}
        with zipfile.ZipFile(self.doc_path) as docx_zip:
            if "word/numbering.xml" not in docx_zip.namelist():
                return num_mapping
            numbering_xml = docx_zip.read("word/numbering.xml")
            numbering_root = parse_xml(numbering_xml)
            numbering_ns = numbering_root.nsmap
            for num in numbering_root.findall("w:num", numbering_ns):
                numId = num.get(qn("w:numId"))
                abstractNumId = num.find("w:abstractNumId", numbering_ns).get(qn("w:val"))
                for abstractNum in numbering_root.findall("w:abstractNum", numbering_ns):
                    if abstractNum.get(qn("w:abstractNumId")) == abstractNumId:
                        for lvl in abstractNum.findall("w:lvl", numbering_ns):
                            ilvl = lvl.get(qn("w:ilvl"))
                            levelText = lvl.find("w:lvlText", numbering_ns).get(qn("w:val"))
                            num_mapping[(numId, ilvl)] = levelText
        return num_mapping

    def detect_title(self):
        titles = []
        for idx, paragraph in enumerate(self.paragraphs):
            if paragraph.get("type") == 'table':
                continue
            # 一级标题识别
            if self._is_first_level_title(paragraph.get("content")):
                titles.append({
                    "index": idx,
                    "content": paragraph.get("content"),
                    "is_first_level": True,
                })
                continue
            # 普通标题识别
            _context = [para.get('content') for para in self.paragraphs[(idx - 2 if idx - 2 >= 0 else 0):idx + 2] if para.get("type") == 'text']
            if self._is_title(paragraph, _context):
                # 若为标题且含有“发行.*?有关机构”
                if bool(re.search(r"发行.*?有关机构", paragraph.get("content"))):
                    break
                titles.append({"index": idx, "content": paragraph.get("content"), "is_first_level": False})
        return titles

    def table_data_format(self, target_tables):
        """
        1. "资产负债表"、"现金流"两个表的数据必须保持一致(募集或者预警通)。
        2. 现金流只需要近两年及一期或者近三年，如果只有近两年的数据，则"资产负债表","现金流"数据全部取预警通。
        3. "有息负债"只取募集，需要注意计算总负债占比时，总负债的统计时间要和有息负债统计时间保持一致
        """
        muji_flag = False
        for pretext_table in target_tables:
            preceding_text = pretext_table.get("pretext")
            tab_datas = pretext_table.get("darray")
            para_elem = pretext_table.get("oxml")
            tab_key = pretext_table.get("name")
            tab_unit = extract_unit(preceding_text)
            if tab_unit and "占比" in str(tab_datas) and "%" not in tab_unit:
                tab_unit = tab_unit + "、%"
            tab_date = extract_date(preceding_text)
            # 获取单位信息
            if tab_unit is None:
                tab_unit = self.get_table_unit(tab_datas)
            # 获取最近的日期
            if tab_date is None:
                tab_date = self.get_recent_date(tab_datas)
            # 处理资产负债表
            if "资产负债表" == tab_key:
                # 保留冗余募集资产负债表
                self.table_datas["资产负债表-募集"] = [{
                    "name": "资产负债表-募集",
                    "date": tab_date,
                    'unit': tab_unit,
                    "preceding_text": preceding_text,
                    "data": tab_datas
                }]
                # 处理表格顺序，单位，金额
                tab_unit, tab_datas = self.asset_liability_data_processing(tab_datas, tab_unit)
                # 防止上下文中单位干扰
                preceding_text = ""

            # 有息负债仅使用募集
            if "有息负债" == tab_key:
                self.table_datas["有息负债-募集"] = [{
                    "name": "有息负债-募集",
                    "date": tab_date,
                    'unit': tab_unit,
                    "preceding_text": preceding_text,
                    "data": tab_datas
                }]
            # 处理现金流表
            if "现金流" == tab_key:
                # 如果表头数小于四（只有近两年数据），则不使用募集，也就不需要进一步处理
                if len(tab_datas[0]) < 4:
                    muji_flag = True
                else:
                    tab_unit, tab_datas = self.cash_flow_data_processing(tab_datas, tab_unit)
            self.table_datas[tab_key] = [{
                "name": tab_key,
                "date": tab_date,
                "unit": tab_unit,
                "preceding_text": preceding_text,
                "data": tab_datas
            }]
            self.table_oxmls[tab_key] = [{
                "name": tab_key,
                "date": tab_date,
                "unit": tab_unit,
                "oxml": para_elem
            }]

        if muji_flag:
            for key in ["资产负债表", "现金流"]:
                self.table_datas.pop(key)

    def cash_flow_data_processing(self, tab_datas, tab_unit):
        """现金流表数据处理"""
        cash_cols = ["经营活动产生的现金流量净额", "投资活动产生的现金流量净额", "筹资活动产生的现金流量净额"]
        try:
            col = tab_datas[0]
            col[0] = "项目"
            assets_df = tab_datas[1:]
            assets_df = pd.DataFrame(assets_df, columns=col)
            assets_df = assets_df.applymap(lambda x: x.strip().replace(" ", "").replace("生产", "产生") if isinstance(x, str) else x)
            assets_df = assets_df[~assets_df.apply(lambda row: all(_cell == "-" for _cell in row[1:]), axis=1)]
            zichan_df = assets_df[assets_df["项目"].isin(cash_cols)].copy()
            zichan_df["项目"] = pd.Categorical(zichan_df["项目"], categories=cash_cols, ordered=True)
            zichan_df_sorted = zichan_df.sort_values(by="项目")
            zichan_df_sorted = zichan_df_sorted[~zichan_df_sorted.apply(lambda row: all(_cell in ["", "-"] for _cell in row[1:]), axis=1)]
            zichan_df_sorted = zichan_df_sorted.dropna(subset=zichan_df_sorted.columns.difference(["项目"]), how='all')
            # 保留近三列数据
            if len(col) > 4 and "年1-" in col[2]:
                zichan_df_sorted = zichan_df_sorted.drop(zichan_df_sorted.columns[2], axis=1)
            elif len(col) > 4 and "年1-" not in col[1]:
                zichan_df_sorted = zichan_df_sorted.drop(zichan_df_sorted.columns[-1], axis=1)
            else:
                zichan_df_sorted = zichan_df_sorted.iloc[:,:4]

            for col in zichan_df_sorted.iloc[:,1:].columns:
                # 处理金额中含有两个点的数据
                mask = zichan_df_sorted[col].astype(str).str.count(r'\.') == 2
                zichan_df_sorted.loc[mask, col] = zichan_df_sorted.loc[mask, col].astype(str).str.replace('.', ',', 1)
                # 处理空数据
                mask_zero = zichan_df_sorted[col].astype(str) == "-"
                zichan_df_sorted.loc[mask_zero, col] = zichan_df_sorted.loc[mask_zero, col].astype(str).str.replace('-', '0.0')
                zichan_df_sorted[col] = zichan_df_sorted[col].str.replace(',', '').str.replace(' ', '').astype("float")
                if tab_unit and "万" in tab_unit:
                    zichan_df_sorted[col] = (zichan_df_sorted[col] / 10000).round(2)
                if tab_unit and "元" == tab_unit:
                    zichan_df_sorted[col] = (zichan_df_sorted[col] / 100000000).round(2)
            zichan_data_list = [zichan_df_sorted.columns.tolist()] + zichan_df_sorted.values.tolist()
            return "亿", zichan_data_list
        except Exception as e:
            logger.exception(e)
            return tab_unit, tab_datas


    def asset_liability_data_processing(self, tab_datas, tab_unit):
        """资产负债表数据处理"""
        zichan_cols = ["货币资金", "结算备付金", "拆出资金", "应收保证金", "应收利息", "应收票据", "应收账款"
                    ,"应收款项融资", "应收保费", "应收分保账款", "应收分保合同准备金", "应收出口退税", "应收补贴款"
                    ,"内部应收款", "预付款项", "其他应收款", "存货", "买入返售金融资产", "交易性金融资产", "衍生金融资产"
                    ,"一年内到期的非流动资产", "待处理流动资产损益", "其他流动资产", "流动资产合计", "发放贷款及垫款", "可供出售金融资产"
                    ,"划分为持有待售的资产", "以公允价值计量且其变动计入其他综合收益的金融资产", "以摊余成本计量的金融资产"
                    ,"债权投资", "其他债权投资", "其他权益工具投资", "其他非流动金融资产", "长期应收款", "长期股权投资", "待摊费用"
                    ,"其他长期投资", "投资性房地产", "固定资产", "合同资产", "在建工程", "使用权资产", "工程物资"
                    ,"生产性生物资产", "公益性生物资产", "油气资产", "无形资产", "开发支出", "商誉", "长期待摊费用"
                    ,"股权分置流通权", "递延所得税资产", "其他非流动资产", "非流动资产合计", "资产总计"]

        fuzhai_cols = ["短期借款", "向中央银行借款", "吸收存款及同业存放", "拆入资金", "交易性金融负债"
            , "衍生金融负债", "卖出回购金融资产款", "应付手续费及佣金", "应付票据", "应付账款", "预收款项"
            , "合同负债", "应付职工薪酬", "应交税费", "应付利息", "应付股利", "其他应交款", "应付保证金"
            , "内部应付款", "其他应付款", "预提费用", "预计流动负债", "应付分保账款", "保险合同准备金"
            , "代理买卖证券款", "代理承销证券款", "国际票证结算", "国内票证结算", "一年内的递延收益", "应付短期债券"
            , "一年内到期的非流动负债", "其他流动负债", "长期借款", "长期应付职工薪酬", "应付债券"
            , "应付债券：优先股", "长期应付款", "预计非流动负债", "长期递延收益"
            , "递延所得税负债", "其他非流动负债", "租赁负债", "担保责任赔偿准备金", "划分为持有待售的负债", "负债合计"]
        zifu_cols = zichan_cols + fuzhai_cols
        try:
            col = tab_datas[0]
            col[0] = "项目"
            assets_df = tab_datas[1:]
            assets_df = pd.DataFrame(assets_df, columns=col)
            assets_df = assets_df.applymap(lambda x: x.strip().replace(" ", "") if isinstance(x, str) else x)
            assets_df = assets_df[~assets_df.apply(lambda row: all(_cell == "-" for _cell in row[1:]), axis=1)]
            zichan_df = assets_df[assets_df["项目"].isin(zifu_cols)].copy()
            zichan_df["项目"] = pd.Categorical(zichan_df["项目"], categories=zifu_cols, ordered=True)
            zichan_df_sorted = zichan_df.sort_values(by="项目")
            zichan_df_sorted["项目"] = zichan_df_sorted["项目"].replace({"流动资产合计": "流动资产", "非流动资产合计": "非流动资产"})
            zichan_df_sorted = zichan_df_sorted.iloc[:, 1:]
            zichan_df_sorted = zichan_df_sorted[~zichan_df_sorted.apply(lambda row: all(_cell in ["", "-"] for _cell in row[1:]), axis=1)]
            zichan_df_sorted = zichan_df_sorted.dropna(subset=zichan_df_sorted.columns.difference(["项目"]), how='all')
            zichan_data_list = None
            s2 = zichan_df.iloc[:, 0]
            for i in range(len(zichan_df_sorted.columns)):
                s1 = zichan_df_sorted.iloc[:, i]
                # 处理金额中含有两个点的数据
                mask = s1.str.count('\.') == 2
                s1[mask] = s1[mask].str.replace('.', ',', 1)
                s1 = s1.str.replace(',', '').str.replace('-', '0').str.replace(' ', '').apply(lambda x: '0' if x == '' else x).astype("float")
                zichan_df_sorted_col = pd.concat([s2, s1], axis=1)
                if tab_unit and "万" in tab_unit:
                    zichan_df_sorted_col.iloc[:, 1] = (zichan_df_sorted_col.iloc[:, 1] / 10000).round(2)
                if tab_unit and "元" == tab_unit:
                    zichan_df_sorted_col.iloc[:, 1] = (zichan_df_sorted_col.iloc[:, 1] / 100000000).round(2)
                zichan_df_sorted_col = zichan_df_sorted_col[zichan_df_sorted_col.iloc[:, 1] >= 1]
                if zichan_df_sorted_col["项目"].str.contains("负债合计", case=False).any():
                    fuzhai_total = zichan_df_sorted_col.loc[zichan_df_sorted_col["项目"] == "负债合计"].iloc[0, 1]
                    self.fuzhai_total.append({"总负债金额": fuzhai_total, "统计时间": zichan_df.columns[i + 1]})
                if i == 0:
                    zichan_data_list = [zichan_df_sorted_col.columns.tolist()] + zichan_df_sorted_col.values.tolist()
            return "亿", zichan_data_list
        except Exception as e:
            logger.exception(e)
            return tab_unit, tab_datas

    def create_table_data_set(self, target_tables):
        table_map = {**{target_table.get("name"): (target_table.get("menu"), target_table.get("pretext"),
                                                   target_table.get("darray")) for target_table in target_tables}}

        for table_type in self.table_type_list:
            if table_map.get(table_type):
                self.table_data_set.append(table_map.get(table_type))
            else:
                self.table_data_set.append([])

    def extract_target_tables_from_doc(self, table_contexts):
        # 生成表格特征
        tableProcessor = TableProcessor(table_contexts, self.private_apply_instance)
        table_feature_datas = tableProcessor.gen_table_feature(self.table_type_list)
        target_tables = tableProcessor.extract_target_tables(self.target_table_descs, table_feature_datas)
        return target_tables

    def filter_table_by_keywords(self, table_contexts):
        start_time = time.time()
        filtered_table_contexts = []
        for table_context in table_contexts:
            # 标题关键词过滤表格
            menu = table_context.get("menu", "")
            # '2. 主营业务-其他(重庆市万州三峡平湖有限公司)'
            if "-其他" in menu:
                continue
            # 上下文关键词过滤表格
            if any(keyword in table_context.get("pretext") for keyword in self.table_keywords) and not any(
                    re.search(keyword, table_context.get("pretext")) for keyword in self.table_keywords_exclude):
                filtered_table_contexts.append(table_context)

        return filtered_table_contexts

    def extract_table_context_with_menu(self):
        start_time = time.time()
        # 表格上下文列表
        table_context_list = []
        # 表格最近段落
        recent_paragraphs = []
        # 当前标题
        current_menu = None
        # 菜单索引
        muidx = 0
        # 遍历段落内容
        for paragraph in self.paragraphs:
            if "text" == paragraph.get("type"):
                text_data = paragraph.get("content")
                # 判断当前文本是否为标题
                if text_data == self.menus[muidx]:
                    muidx = muidx + 1
                    current_menu = text_data
                    # 重置表格上下文
                    recent_paragraphs = []
                if muidx + 1 >= len(self.menus):
                    muidx = len(self.menus) - 1
                # 排除页眉页脚
                if bool(re.search(r'^[0-9]+$', text_data)) or bool(re.search(r'.{6,}募集说明书$', text_data)):
                    continue
                # 排除多个单位
                if text_data.startswith("单位") and recent_paragraphs and recent_paragraphs[-1].startswith("单位"):
                    continue
                # 添加段落文本到 recent_paragraphs , preceding_length=2
                recent_paragraphs.append(text_data)
                if len(recent_paragraphs) > 2:
                    recent_paragraphs.pop(0)
            elif "table" == paragraph.get("type"):
                if not current_menu:
                    continue
                table_data = paragraph.get("content")
                para_elem = paragraph.get("oxml")
                # 如果出现连续表格，说明属于同一张表，合并数据
                if table_context_list and not recent_paragraphs:
                    befare_table = table_context_list[-1]
                    befare_table["darray"].extend(table_data)
                    befare_table["content"] = "\n".join(["\t".join(row) for row in befare_table["darray"]][0:40])
                    rows = para_elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
                    for row in rows:
                        befare_table["oxml"].append(row)
                    continue
                table_context_list.append({
                    "id": str(uuid.uuid4()),
                    "menu": current_menu,
                    "pretext": "\n".join(recent_paragraphs),
                    "content": "\n".join(["\t".join(row) for row in table_data][0:40]),
                    "darray": table_data,
                    "oxml": para_elem
                })
                recent_paragraphs = []
        return table_context_list

    def get_recent_date(self, table_datas):
        for row_data in table_datas:
            for cell_data in row_data:
                tmp_date = extract_date(cell_data)
                if tmp_date is not None:
                    return tmp_date

    def get_table_unit(self, table_datas):
        for row_data in table_datas:
            for cell_data in row_data:
                tmp_unit = extract_unit(cell_data)
                if tmp_unit is not None:
                    return tmp_unit

if __name__ == "__main__":
    root_path = "D:/opt/城投债/20250102/company_data/盐城市城镇化建设投资集团有限公司"
    doc_file_name = "1-2 盐城市城镇化建设投资集团有限公司2024年面向专业投资者非公开发行公司债券(第二期)募集说明书.docx"
    pdf_file_name = "1-2 盐城市城镇化建设投资集团有限公司2024年面向专业投资者非公开发行公司债券(第二期)募集说明书.pdf"
    prospectus_doc_path = os.path.join(root_path, doc_file_name)
    prospectusDoc = SplitDoc(prospectus_doc_path, "")
    logger.debug("======================================")
    for key, value in prospectusDoc.table_datas.items():
        logger.debug("**************************************")
        logger.debug(key)
        logger.debug(value)
