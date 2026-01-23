from docx import Document
from utils.file_util import *
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from utils.log_utils import get_logger

logger = get_logger()


class OrganizeDoc():
    def __init__(self, src_dir):
        self.src_dir = src_dir

    def instr(self, string, list):
        if any(x in string for x in list):
            return True
        else:
            return False

    def parse_oxml_to_2list(self, oxml_table, namespaces):
        """将表格数据转为二维数组数据"""
        curr_table = []
        for row in oxml_table.iterfind("w:tr", namespaces):
            cell_idx = -1
            row_data = []
            for cell in row.iterfind("w:tc", namespaces):
                cell_idx += 1
                vmerge = True if cell.find('.//w:vMerge', namespaces) is not None else False
                text = "".join([para.text for para in cell.iterfind("w:p", namespaces)])
                if vmerge and text == "":
                    text = curr_table[-1][cell_idx]
                # 获取列合并
                grid_span = cell.find('.//w:gridSpan', namespaces)
                if grid_span is not None:
                    grid_span_val = grid_span.val
                    row_data += [text for _ in range(grid_span_val)]
                else:
                    row_data.append(text)
            curr_table.append(row_data)
        return curr_table

    def get_company_name(self, file_path):
        logger.info(file_path)
        _company_name = None
        _doc = Document(file_path)
        _doc_elem = parse_xml(_doc.part.element.xml)
        _body_elem = _doc_elem.find("w:body", _doc_elem.nsmap)
        com_keyword = r"(.*有限公司)|(.*投资公司)|(.*有限责任公司)|(.*投资经营公司)|(.*总公司)"
        for para_elem in _body_elem.getchildren():
            if _company_name:
                break
            if para_elem.tag == qn("w:p"):
                text = para_elem.text.strip()
                match = re.search(com_keyword,
                                  para_elem.text.strip())
                if match:
                    _company_name = match.group(0)
                if "募集说明书" in para_elem.text:
                    break
            elif para_elem.tag == qn("w:tbl"):
                _row_list = self.parse_oxml_to_2list(para_elem, _doc_elem.nsmap)
                for _row in _row_list:
                    if _company_name:
                        break
                    for _cell in _row:
                        match = re.search(com_keyword, _cell)
                        if match:
                            _company_name = match.group(0).strip()
                            break
        if _company_name is None:
            logger.error(f"无公司名称,募集说明书路径:{file_path}")

        return _company_name

    def which_company_name(self, file_path, company_names):
        _doc = Document(file_path)
        for _para in _doc.paragraphs:
            if len(_para.text) < 25:
                continue
            for company_name in company_names:
                if company_name in _para.text:
                    return company_name

    def organize(self):
        company_name = None
        try:
            # TODO 增加判断是否有多个募集说明书，有则抛出异常
            for root, dirs, files in os.walk(self.src_dir):
                for file in files:
                    if self.instr(file, ['募集说明书', '调查报告', '尽调报告', "授信报告"]) and file.endswith(
                            tuple([".docx", ".doc"])) and not file.startswith("~$"):
                        _file = file
                        _src_dir = os.path.join(root, file)
                        if file.endswith(".doc"):
                            convert_doc_to_docx(_src_dir, _src_dir + "x")
                            logger.debug(f"doc to docx succeed. file: {_file}")
                            _file = _file + "x"
                            _src_dir = _src_dir + "x"
                        if "募集说明书" in file:
                            company_name = self.get_company_name(_src_dir)
                            logger.info(f"公司名称：{company_name}")
                            # 募集转为pdf
                            # convert_doc_to_pdf(_src_dir, os.path.join(self.src_dir, f"{os.path.splitext(file)[0]}.pdf"))
                            # logger.debug(f"doc to pdf succeed. file: {_file}")
            if not company_name:
                raise Exception("文件不符合要求。")
            return company_name
        except Exception as e:
            logger.exception(e)
            raise Exception(f"文件转换错误。{str(e)}")


if __name__ == "__main__":
    src_dir = "D:\\data\\20250304"
    od = OrganizeDoc(src_dir)
    print(od.organize())
